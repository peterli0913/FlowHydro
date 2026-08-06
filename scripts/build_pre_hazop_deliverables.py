"""
Build three deliverables based on Keith's Risk Register (T-2502 v0.1)
plus internal alignment work:

1. Update Meeting Minutes Tracking Record 0701.xlsx (add missing
   Outstanding Questions + all 59 Risk Register items to sheet '0701')
2. Create bilingual Word version of the Risk Register
   (Risk_Register_Bilingual.docx) — matches the style of
   Outstanding_Questions_for_Keith.docx
3. Create internal priority warning document
   (Internal_Priority_Warning_Risks.docx) — for the 5-6 risks the
   Chinese team has NOT yet fully considered
"""

import copy
import os
import shutil
from datetime import datetime

# -----------------------------------------------------------------
# Data model
# -----------------------------------------------------------------
# All 59 risks from T-2502 Risk Register v0.1 with bilingual content.
# Each item: (id, category_en, category_cn, desc_en, desc_cn,
#             response_en, response_cn, owner, closed, related_oq,
#             is_new_awareness)

RISKS = [
    # ------- Engineering Design -------
    ("R1", "Engineering Design", "工程设计",
     "Design Temperature — Vessel could fail during relief event. Equipment design temperature is specified for process requirements and does not consider elevated temperatures during relief events when short-term temperature increase has the potential to exceed the design temperature.",
     "设计温度 —— 容器可能在泄放事件中失效。设备设计温度是按工艺需求设定的，未考虑泄放事件中的短时温升可能超过设计温度的情形。",
     "", "",
     "Ziliang Zhao/Chaoqun Hu",
     False, "Q4.1", False),

    ("R2", "Engineering Design", "工程设计",
     "Design Requirements — L:G ratio for plant is not clearly defined. Equipment design could result in very limited operating ranges and incorrect assumptions in plant capability and relief design.",
     "设计条件 —— 装置的液气比（L:G ratio）未明确定义。可能导致设备设计的操作区间受限，以及在装置能力和泄放设计中做出错误假设。",
     "", "",
     "Shaofeng Gao/Chaoqun Hu",
     False, "", True),

    ("R3", "Engineering Design", "工程设计",
     "Design Requirements — Catalyst specification is not available. Design may not be suitable for some catalysts. Dusty catalysts could cause blockages.",
     "设计条件 —— 催化剂规格未明确。设计可能不适用于某些催化剂；含粉尘的催化剂可能导致堵塞。",
     "Design philosophy is based on catalyst beads/pellets which are dust-free, non-flammable and have exposure levels requiring OEB 3 or lower containment. Note: OEB 3 containment is a deviation from the URS and requires acceptance.",
     "设计基于无粉尘、非燃、暴露等级 OEB 3 或更低的颗粒/粒状催化剂。注：OEB 3 的密闭要求相对 URS 存在偏离，需要接受确认。",
     "Chaoqun Hu",
     False, "", False),

    ("R4", "Engineering Design", "工程设计",
     "Access — Layouts are not available to assess whether access to equipment is acceptable for operations and maintenance.",
     "通道 —— 布置图尚未提供，无法评估设备的操作与维护通道是否满足要求。",
     "", "",
     "Ziliang Zhao/All",
     False, "Q2.1", False),

    ("R5", "Engineering Design", "工程设计",
     "Maintenance — Lifting Aids (davit arms, lifting beams, etc.) have not been considered in design, with potential for injury to maintenance personnel removing heavy equipment for servicing / repair / replacement.",
     "维护 —— 设计中尚未考虑起重辅助设施（吊臂、起吊梁等），维护人员在拆卸重型设备进行检修 / 修理 / 更换时可能存在人身伤害风险。",
     "", "",
     "Ziliang Zhao/Shuangshuang Fan",
     False, "", True),

    ("R6", "Engineering Design", "工程设计",
     "Lighting — Layouts are not available to assess whether lighting for routine operations and maintenance activities is acceptable.",
     "照明 —— 布置图尚未提供，无法评估日常操作和维护活动的照明条件是否满足要求。",
     "", "",
     "Ziliang Zhao/All",
     False, "", True),

    # ------- Continuous reactor build -------
    ("R7", "Continuous Reactor Build", "连续反应器建造",
     "System leak detection — Pressure testing of high-pressure system will be hydraulically tested. Hydraulic testing demonstrates mechanical strength of the pressure system but does not demonstrate integrity for a hydrogen system.",
     "系统检漏 —— 高压系统将进行水压试验。水压试验只能证明压力系统的机械强度，无法证明氢气系统的密封完整性。",
     "Helium testing is not routinely carried out in China.",
     "国内通常不常规进行氦气检漏。",
     "Ziliang Zhao/Shuangshuang Fan",
     False, "", True),

    ("R8", "Continuous Reactor Build", "连续反应器建造",
     "Line sizing / specification — Vents / headers incorrectly sized / specified, resulting in high back pressure and safety / operability issues, e.g. two-phase flow in emergency vent relief header.",
     "管线规格 / 尺寸 —— 放空 / 泄放总管尺寸或规格不正确，导致背压过高及安全 / 可操作性问题，例如紧急泄放总管中出现两相流。",
     "", "",
     "Chaoqun Hu",
     False, "", True),

    ("R9", "Continuous Reactor Build", "连续反应器建造",
     "Insulation — Insulation requirements not specified, which could result in equipment / pipe layout not allowing space for insulation to be installed.",
     "保温 —— 保温需求尚未明确，可能导致设备 / 管路布置未预留保温安装空间。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "", True),

    ("R10", "Continuous Reactor Build", "连续反应器建造",
     "Modular design — Modular design could introduce significant increase in site construction activities, e.g. mechanical and EC&I installation and subsequent testing.",
     "模块化设计 —— 模块化设计可能大幅增加现场施工工作量，例如机械与电控仪表安装及后续测试。",
     "", "",
     "Ziliang Zhao",
     False, "", False),

    ("R11", "Continuous Reactor Build", "连续反应器建造",
     "Skid assembled on site could result in CE / UKCA Responsible Person / Authorised Representative in China not being able to issue Declaration of Conformity for Skid Package.",
     "撬块在现场组装可能导致中国境内的 CE / UKCA 负责人 / 授权代表无法为撬块整机签发符合性声明（Declaration of Conformity）。",
     "", "",
     "Ziliang Zhao",
     False, "", True),

    ("R12", "Continuous Reactor Build", "连续反应器建造",
     "Wall penetrations — Layouts and cable schedules not available to confirm space is available for pipe / cable penetrations to / from R23 Module.",
     "穿墙孔 —— 布置图和电缆清单尚未提供，无法确认与 R23 模块之间的管道 / 电缆穿墙位置有足够空间。",
     "", "",
     "Dezhi Meng",
     False, "Q2.6", False),

    # ------- Catalyst Charging -------
    ("R13", "Catalyst Charging", "催化剂加料",
     "Catalyst charge point — Inlet pipe diameter is small and blockages may occur.",
     "加料口 —— 加料管径较小，可能发生堵塞。",
     "Inlet connection maximised (DN32 for CR01 and DN25 for CR02/CR03). Material is deemed to be free-flowing beads or pellets with no risk of blocking inlet pipe.",
     "加料接口已最大化（CR01 采用 DN32，CR02/CR03 采用 DN25）。物料按自由流动的颗粒 / 珠状催化剂考虑，无堵塞加料管的风险。",
     "Shuangshuang Fan",
     True, "", False),

    ("R14", "Catalyst Charging", "催化剂加料",
     "Catalyst physical property requirements (size, shape, packing, etc.) are not defined, with the potential for incompatible catalyst to be charged to the reactors.",
     "催化剂物性要求（粒径、形状、装填密度等）尚未定义，存在向反应器加入不兼容催化剂的风险。",
     "", "",
     "Chaoqun Hu",
     False, "", False),

    ("R15", "Catalyst Charging", "催化剂加料",
     "Packing of catalyst is not controlled, with the risk of reactants / gas channelling in the continuous reactor column, and the potential for incomplete reaction or impurities forming.",
     "催化剂装填未受控，存在反应物 / 气体在连续反应柱内形成沟流的风险，可能导致反应不完全或杂质产生。",
     "", "",
     "Chaoqun Hu",
     False, "", False),

    ("R16", "Catalyst Charging", "催化剂加料",
     "Contained charge system design not available to assess risks and / or design limitation.",
     "密闭加料系统的设计尚未提供，无法评估其风险与设计局限。",
     "", "",
     "Shuangshuang Fan",
     False, "", True),

    ("R17", "Catalyst Charging", "催化剂加料",
     "Manual handling requirements for catalyst charging not assessed.",
     "催化剂加料的人工搬运需求尚未评估。",
     "", "",
     "Shuangshuang Fan",
     False, "", False),

    ("R18", "Catalyst Charging", "催化剂加料",
     "Access to charge point and support of charge container / bag not designed (access steps / integrated support shelf).",
     "加料口的通道以及加料容器 / 料袋的支撑尚未设计（例如操作台阶、集成支撑架）。",
     "", "",
     "Shuangshuang Fan",
     False, "", False),

    # ------- Reactor Feed -------
    ("R19", "Reactor Feed", "反应器进料",
     "20-micron catalyst retaining screen — Screen blockages prevent / limit flow of gas / reaction mix to continuous reactors.",
     "20 微米催化剂拦截筛 —— 拦截筛堵塞会阻止 / 限制气体 / 反应液流入连续反应器。",
     "China Operations have reported blockages occur in local plant and screen has been removed and discarded.",
     "国内运行反馈已在本地装置中发生堵塞，已将拦截筛拆除并弃用。",
     "Chaoqun Hu",
     True, "", False),

    ("R20", "Reactor Feed", "反应器进料",
     "Polishing filter (FTR9832) — Filter blockages prevent / limit flow of gas / reaction mix through continuous reactors. Pressure increases in reactor system, with potential for undesirable reactions to occur, until reactor high-pressure trip is activated.",
     "精滤器 (FTR9832) —— 滤芯堵塞会阻止 / 限制气体 / 反应液流经连续反应器；反应器压力升高，直到触发高压跳停前可能发生副反应。",
     "", "",
     "Chaoqun Hu",
     False, "", False),

    # ------- Reactors -------
    ("R21", "CR01 – CR03 Reactors", "CR01 – CR03 反应器",
     "Temperature control — Jacket temperature control is configured to control jacket temperature only. Reactor internal temperatures are monitored, but no control, alarms or interlocks are configured for process control.",
     "温度控制 —— 夹套温度仅控制夹套温度本身。反应器内部温度虽被监测，但未针对工艺过程配置控制、报警或联锁。",
     "Scale-up provides large operating envelopes for process temperature requirements and additional control, alarms and interlocks are not required.",
     "由于放大试验提供了较大的工艺温度操作裕度，无需增加额外的控制、报警或联锁。",
     "Chaoqun Hu",
     True, "", False),

    # ------- System Pressure -------
    ("R22", "System Pressure Monitoring", "系统压力监测",
     "Differential pressure monitoring — Accuracy of static pressure measurements may not enable differential pressures to be calculated with any degree of accuracy or certainty.",
     "差压监测 —— 静态压力测量精度可能不足以支持差压计算达到所需的精度或置信度。",
     "China operations have reported blockages occur in local plant and screen has been removed and discarded.",
     "国内运行反馈本地装置中已发生堵塞，已将拦截筛拆除并弃用。",
     "Yu Gao",
     False, "", False),

    # ------- Over-pressure -------
    ("R23", "Over Pressure Protection", "过压保护",
     "Relief system design — Pressure systems not identified and relief studies not available. Identification of worst-case relief scenario not identified, and relief streams may be incorrectly categorised. Details of maximum reaction exotherm not fully documented.",
     "泄放系统设计 —— 压力系统尚未识别、泄放研究尚未提供。最坏工况泄放情景未识别，泄放物流可能被错误归类。反应最大放热的详细信息尚未完整记录。",
     "", "",
     "Chaoqun Hu/Shaofeng Gao",
     False, "Q4.1", False),

    # ------- Separator -------
    ("R24", "Separator", "气液分离器",
     "Vessel design — Complete degassing of product stream not achieved, with the potential for hydrogen carry-over into BT04 and product transfer containers.",
     "容器设计 —— 产品物流未实现完全脱气，H₂ 可能带入 BT04 及产品转运容器。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "", True),

    ("R25", "Separator", "气液分离器",
     "Cleaning of separator not fully understood but current concept relies on vessel free-draining to BT04. Current spray ball sizing is not consistent with this approach and may require significant rework to process piping.",
     "分离器清洗方案尚未完全明确，当前概念依赖容器向 BT04 的自重排放。现有喷淋球规格与该思路不匹配，可能需要对工艺管路进行较大改动。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "Q1.4,Q1.5", False),

    ("R26", "Separator", "气液分离器",
     "Cleaning of LIT003 housing does not appear to be considered in design.",
     "LIT003 液位变送器外壳的清洗在设计中似乎尚未考虑。",
     "", "",
     "Yu Gao/Chaoqun Hu",
     False, "Q1.7", False),

    ("R27", "Separator", "气液分离器",
     "P&ID does not clearly show nozzle locations and could lead to mal-ops / misunderstanding of the design intent by the Operations team.",
     "P&ID 上未清晰标注管嘴位置，可能导致运行团队误操作或误解设计意图。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "Q2.3", False),

    # ------- Sampling -------
    ("R28", "Sampling", "取样系统",
     "Sampler design — When depressurising the sample collection chamber, high gas or vapour vent rates could result in a bubbly or churn-turbulent flow regime, and sample will discharge to vent line as two-phase flow.",
     "取样器设计 —— 对取样收集腔进行减压时，气体或蒸汽的高排放速率可能形成鼓泡或搅动湍流流态，样品将以两相流形式排入放空管。",
     "", "",
     "Chaoqun Hu",
     False, "Q5.1", False),

    ("R29", "Sampling", "取样系统",
     "With significantly more gas being present in the process stream, what will prevent overpressure of sample bottle and low-pressure nitrogen pipework?",
     "由于工艺物流中气体明显较多，采用什么措施来防止取样瓶和低压氮气管路的超压？",
     "", "",
     "Chaoqun Hu/Yu Gao",
     False, "", True),

    # ------- Catalyst Discharging -------
    ("R30", "Catalyst Discharging", "催化剂卸料",
     "Catalyst Discharge Plug valve — Outlet pipe diameter is small and blockages may occur.",
     "催化剂卸料闸阀 —— 出口管径较小，可能发生堵塞。",
     "Outlet connection maximised (DN32 for CR01 and DN25 for CR02/CR03). Material is deemed to be free-flowing beads or pellets with no risk of blocking outlet pipe.",
     "出口接管已最大化（CR01 采用 DN32，CR02/CR03 采用 DN25）。物料按自由流动的颗粒 / 珠状催化剂考虑，无堵塞出口管的风险。",
     "Shuangshuang Fan",
     True, "", False),

    ("R31", "Catalyst Discharging", "催化剂卸料",
     "Contained discharge system design not available. Layout design may not provide sufficient headroom for contained discharge system.",
     "密闭卸料系统的设计尚未提供。当前布置可能无法为密闭卸料系统提供足够的顶部空间。",
     "", "",
     "Shuangshuang Fan",
     False, "", True),

    ("R32", "Catalyst Discharging", "催化剂卸料",
     "Complete removal of catalyst during discharging operations is unlikely to remove 100% of the material, with the risk of cross contamination of catalyst material between campaigns.",
     "卸料操作难以完全（100%）清除催化剂，存在批次间催化剂交叉污染的风险。",
     "", "",
     "Shuangshuang Fan/Chaoqun Hu",
     False, "", False),

    # ------- Cleaning -------
    ("R33", "Cleaning", "清洗",
     "CIP Design — CIP design incomplete. Initial spray ball sizing is inconsistent with process line sizing approach and may require significant rework to process piping and/or equipment layout.",
     "CIP 设计 —— 就地清洗设计不完整。初步喷淋球规格与工艺管路选型方式不匹配，可能需对工艺管路和 / 或设备布置进行较大改动。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "Q1.1,Q1.5", False),

    ("R34", "Cleaning", "清洗",
     "Cleaning carried out at nominally atmospheric pressure with the risk of material being trapped between flanges, which tend to separate slightly during high-pressure operation, and could lead to cross contamination between batches.",
     "清洗在近常压下进行，高压运行时法兰会出现微量分开，可能将物料残留于法兰间隙中，进而导致批次间交叉污染。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "", True),

    ("R35", "Cleaning", "清洗",
     "Validation / Verification — Cleaning validation / verification philosophy not available. QA may not accept that the continuous reactor achieves minimum cleaning standards.",
     "验证 / 确认 —— 清洗验证 / 确认策略尚未提供。质量保证部门可能不接受连续反应器满足最低清洗标准。",
     "", "",
     "Chaoqun Hu/Shaofeng Gao",
     False, "Q1.2,Q1.9", False),

    ("R36", "Cleaning", "清洗",
     "Layout — Layouts are not available to evaluate the impact of bypasses, dead legs, drains, vents, etc.",
     "布置 —— 布置图尚未提供，无法评估旁路、死管、排放、放空等的影响。",
     "", "",
     "Ziliang Zhao/Shuangshuang Fan",
     False, "", True),

    ("R37", "Cleaning", "清洗",
     "Layout — Layouts are not available to evaluate access to equipment for visual inspections, which may impact the effectiveness of critical inspections.",
     "布置 —— 布置图尚未提供，无法评估设备的目视检查通道，可能影响关键检查的有效性。",
     "", "",
     "Ziliang Zhao/Shuangshuang Fan",
     False, "", False),

    ("R38", "Cleaning", "清洗",
     "Layout — Equipment design / layout does not allow for routine disassembly and manual cleaning. Equipment turnaround times and operating costs would be adversely impacted if CIP system cannot be validated.",
     "布置 —— 设备设计 / 布置未考虑日常拆装和人工清洗。如 CIP 系统无法通过验证，设备周转时间和运行成本都会受到不利影响。",
     "", "",
     "Ziliang Zhao/Shuangshuang Fan",
     False, "", False),

    # ------- FT001 -------
    ("R39", "FT001", "FT001 流量计",
     "Liquid Stream Flow Sensor for flow control loop FIC001. FT001 is a Coriolis mass flow meter installed downstream of the metering pump MP01. Pulsed flow from the pump will impact on the flow measurement and accurate flow control may not be possible.",
     "液相流量传感器（用于 FIC001 流控回路）—— FT001 是安装在计量泵 MP01 下游的科里奥利质量流量计。泵的脉冲流会影响流量测量，可能无法实现精确流量控制。",
     "China Team have previously used this configuration and not experienced flow control issues. No further action required.",
     "国内团队此前采用过该配置，未出现流控问题。无需进一步动作。",
     "Yu Gao",
     True, "", False),

    # ------- MP01 -------
    ("R40", "MP01 Metering Pump", "MP01 计量泵",
     "Maximum flow may exceed separator design limitations and carry hydrogen over into BT04 / product collection containers.",
     "泵的最大流量可能超出分离器的设计上限，将 H₂ 带入 BT04 / 产品收集容器。",
     "", "",
     "Chaoqun Hu/Shuangshuang Fan",
     False, "", False),

    ("R41", "MP01 Metering Pump", "MP01 计量泵",
     "Pump turndown does not meet URS minimum flow requirements.",
     "泵的调节比无法满足 URS 中规定的最低流量要求。",
     "", "",
     "Chaoqun Hu",
     False, "", False),

    # ------- FIC003 -------
    ("R42", "FIC003", "FIC003 氢气流量控制",
     "Hydrogen gas flow controller — Flow controller does not support tight shut-off and system pressure increase during hold periods.",
     "氢气流量控制器 —— 控制器不支持严密关断，系统在保压期间压力会持续上升。",
     "Additional shut-off valve to be installed downstream of the controller to provide tight shut-off.",
     "在控制器下游追加一台切断阀以提供严密关断。",
     "Yu Gao",
     True, "", False),

    ("R43", "FIC003", "FIC003 氢气流量控制",
     "Maximum flow through controller may exceed separator design limitations and carry liquid into the vent line.",
     "流经控制器的最大流量可能超出分离器设计上限，将液体带入放空管。",
     "High level interlock installed in the separator to stop gas supply and close PIC028 shut-off valve (CV08).",
     "在分离器上设置高液位联锁，停 H₂ 供给并关闭 PIC028 的切断阀（CV08）。",
     "Yu Gao/Chaoqun Hu",
     True, "", False),

    ("R44", "FIC003", "FIC003 氢气流量控制",
     "Time required for pressure testing the system and nitrogen inertion may be excessively high.",
     "系统加压试验及氮气钝化所需时间可能过长。",
     "Pressure test / inertion times to be evaluated during FIC003 selection. Manual vent installed on BT03 to reduce depressurisation times via FIC003.",
     "在 FIC003 选型时评估压力试验 / 惰化时间。在 BT03 上加装人工放空阀，缩短经 FIC003 的减压时间。",
     "Yu Gao/Chaoqun Hu",
     True, "", False),

    ("R45", "FIC003", "FIC003 氢气流量控制",
     "Differential pressure across the controller exceeds manufacturer's recommendations and controller may not function correctly.",
     "控制器两端的差压超出厂家建议范围，可能导致控制器无法正常工作。",
     "China Team reviewed with manufacturer and accuracy of flow control may be compromised but this is not deemed significant and no further action is required.",
     "国内团队已与厂家复核，流控精度可能有所下降但不显著，无需进一步动作。",
     "Yu Gao",
     True, "", False),

    # ------- LIC003 -------
    ("R46", "LIC003", "LIC003 分离器液位",
     "SE01 Separator Level Indicator — Cleaning / decontamination of housing is not possible. Residual product in housing could lead to cross contamination between campaigns.",
     "SE01 分离器液位指示 —— 变送器外壳无法进行清洗 / 净化。外壳内的残留物料可能导致批次间交叉污染。",
     "Design relies on dissolution of residual material. 100% flood-fill of housing is not possible, therefore any residual trace contamination in the housing is deemed to be small and will not significantly impact cleaning verification.",
     "设计依赖残余物料的溶解。外壳无法 100% 灌满，因此外壳内任何残余痕量污染被视为很小，不会显著影响清洗验证。",
     "Yu Gao/Chaoqun Hu",
     True, "", False),

    # ------- FIC004 -------
    ("R47", "FIC004", "FIC004 分离器液位控制",
     "Separator outlet flow / level control — Differential pressure across the controller exceeds manufacturer's recommendations and controller capability.",
     "分离器出口流量 / 液位控制 —— 控制器两端差压超出厂家建议范围及控制器能力。",
     "", "",
     "Yu Gao/Chaoqun Hu",
     False, "", False),

    # ------- FIC005 -------
    ("R48", "FIC005", "FIC005 BT01 氮气吹扫",
     "BT01 nitrogen sweep flow controller — Differential pressure across the controller is less than manufacturer's recommendations and controller may not function correctly.",
     "BT01 氮气吹扫流量控制器 —— 控制器两端差压低于厂家建议范围，可能导致控制器无法正常工作。",
     "China Team reviewed with manufacturer and accuracy of flow control may be compromised but this is not deemed significant and no further action is required.",
     "国内团队已与厂家复核，流控精度可能有所下降但不显著，无需进一步动作。",
     "Yu Gao",
     True, "", False),

    # ------- FIC007 -------
    ("R49", "FIC007", "FIC007 BT04 氮气吹扫",
     "BT04 nitrogen sweep flow controller — Differential pressure across the controller is less than manufacturer's recommendations and controller may not function correctly.",
     "BT04 氮气吹扫流量控制器 —— 控制器两端差压低于厂家建议范围，可能导致控制器无法正常工作。",
     "China Team reviewed with manufacturer and accuracy of flow control may be compromised but this is not deemed significant and no further action is required.",
     "国内团队已与厂家复核，流控精度可能有所下降但不显著，无需进一步动作。",
     "Yu Gao",
     True, "", False),

    # ------- PIC028 -------
    ("R50", "PIC028", "PIC028 系统背压控制",
     "System back pressure controller — The back pressure controller does not support tight shut-off during hold periods and system pressure will not be maintained.",
     "系统背压控制器 —— 该控制器在保压期间不支持严密关断，系统压力无法保持。",
     "Additional shut-off valve to be installed downstream of the controller.",
     "在该控制器下游追加一台切断阀。",
     "Yu Gao",
     True, "", False),

    ("R51", "PIC028", "PIC028 系统背压控制",
     "Time required for depressurising the system following pressure testing and nitrogen inertion may be excessively high.",
     "加压试验和氮气钝化之后系统减压所需时间可能过长。",
     "Manual bypass installed around PIC028 controller.",
     "在 PIC028 控制器周围加装人工旁路。",
     "Yu Gao",
     True, "", False),

    ("R52", "PIC028", "PIC028 系统背压控制",
     "Differential pressure across the controller exceeds manufacturer's recommendations and controller may not function correctly.",
     "控制器两端差压超出厂家建议范围，可能导致控制器无法正常工作。",
     "", "",
     "Yu Gao",
     False, "", False),

    # ------- PIC030 -------
    ("R53", "PIC030", "PIC030 分离器液相出口减压",
     "Separator liquid outlet pressure reducer — Is this confirmed as a Bronkhorst pressure controller, or is a mechanical regulator still an option? Does this provide the required independence assumed in the relevant SIL / risk assessments?",
     "分离器液相出口减压器 —— 该阀是否确认为 Bronkhorst 电子压力控制器？还是仍可能选用机械式减压阀？该配置能否提供相关 SIL / 风险评估中假设所需的独立性？",
     "", "",
     "Yu Gao/Chaoqun Hu",
     False, "", False),

    # ------- PIC049 -------
    ("R54", "PIC049", "PIC049 BT04 背压控制",
     "BT04 back pressure controller — Differential pressure across the controller is less than manufacturer's recommendations and controller may not function correctly.",
     "BT04 背压控制器 —— 控制器两端差压低于厂家建议范围，可能导致控制器无法正常工作。",
     "China Team reviewed with manufacturer and accuracy of flow control may be compromised but this is not deemed significant and no further action is required.",
     "国内团队已与厂家复核，流控精度可能有所下降但不显著，无需进一步动作。",
     "Yu Gao",
     True, "", False),

    # ------- Automation -------
    ("R55", "Automation", "自动化",
     "Automation Build — Philosophy and Design documentation not available to assess risks and limitations.",
     "自动化建设 —— 自动化策略与设计文档尚未提供，无法评估其风险与局限。",
     "", "",
     "Dezhi Meng",
     False, "Q3.1,Q3.2", False),

    # ------- Documentation -------
    ("R56", "Documentation", "文档",
     "Technical File — Technical File not available to assess risks and limitations.",
     "技术文件 —— 技术文件尚未提供，无法评估其风险与局限。",
     "", "",
     "Ziliang Zhao/Shaofeng Gao",
     False, "Q4.5,Q7.1", False),

    ("R57", "Documentation", "文档",
     "O&M Manual — O&M manual not available to assess risks and limitations.",
     "操作维护手册 —— O&M 手册尚未提供，无法评估其风险与局限。",
     "", "",
     "Shaofeng Gao/All",
     False, "Q7.1", False),

    ("R58", "Documentation", "文档",
     "Emissions Calculations — Calculations to demonstrate expected release of Class A and Class B VOCs to atmosphere are not available. These are required to demonstrate compliance with EA Permit limits.",
     "排放量计算 —— 尚未提供证明 A 类和 B 类 VOC 向大气排放的预期计算。该计算是证明满足 EA 许可（英国环境署许可）限值要求所必需。",
     "", "",
     "Chaoqun Hu/Shaofeng Gao",
     False, "", True),

    # ------- Regulatory compliance -------
    ("R59", "Regulatory Compliance", "法规合规",
     "CE | Functional Safety | DSEAR / ATEX | PSSR | COSHH | PUWER (Access, manual handling) | Working at Height (Access Steps / Support Shelf) | LOLER (Davit arms) | Lighting | Occupational Hygiene | EA Permit — Regulatory non-conformance would prevent beneficial use of equipment in the UK.",
     "CE 认证 | 功能安全 | DSEAR / ATEX 防爆 | PSSR 压力系统安全 | COSHH 危害物质 | PUWER 工作设备（含通道与人工搬运）| 高处作业（操作台阶 / 支撑架）| LOLER 起重设备（吊臂）| 照明 | 职业卫生 | EA 环境许可 —— 任何法规合规缺失都会导致设备无法在英国正常投入使用。",
     "", "",
     "All/Peter",
     False, "", True),
]


# -----------------------------------------------------------------
# 3 missing Outstanding Questions (Q1.9, Q7.2, Q7.3)
# -----------------------------------------------------------------
MISSING_OQ = [
    {
        "section_en": "Cleaning Philosophy",
        "section_cn": "清洗策略与清洗程序",
        "title_en": "Cleaning sign-off authority and record",
        "title_cn": "清洗放行权限与记录",
        "body_en": "Please confirm the role authorised to approve cleaning completion prior to the next campaign, and the documentation and records required to support this sign-off (batch record, cleaning log, sample analysis certificate, etc.).",
        "body_cn": "请确认在下一批次前有权审批清洗完成的岗位，以及支持该放行签批的文档与记录要求（批记录 / 清洗日志 / 样品分析证书等）。",
        "owner": "Chaoqun Hu/Shaofeng Gao",
    },
    {
        "section_en": "Documentation Status",
        "section_cn": "文档状态",
        "title_en": "Pre-HAZOP document review period",
        "title_cn": "HAZOP 前文档评审期",
        "body_en": "Please advise the schedule for the pre-HAZOP document review period, during which the HAZOP study team will formally review the full document set prior to the HAZOP session. A minimum two-week review period ahead of the planned mid-August HAZOP is requested.",
        "body_cn": "请说明 HAZOP 前文档评审期的时间安排。HAZOP 分析小组将在本期内对完整文件集进行正式评审。请在 8 月中旬拟定 HAZOP 前预留至少两周的评审时间。",
        "owner": "Peter/Shaofeng Gao",
    },
    {
        "section_en": "Documentation Status",
        "section_cn": "文档状态",
        "title_en": "HAZOP session logistics",
        "title_cn": "HAZOP 会议后勤安排",
        "body_en": "Please confirm the target HAZOP session dates (mid-August 2026), the appointed HAZOP chair, the study team composition (site operations representative, process engineer, control & safety engineer, mechanical engineer, and IEPE / CFCT design representatives), the location, and any pre-work required from the design team.",
        "body_cn": "请确认 HAZOP 会议目标日期（2026 年 8 月中旬）、拟任 HAZOP 主席、分析小组构成（现场操作代表、工艺工程师、控制与安全工程师、机械工程师、以及 IEPE / CFCT 设计代表）、会议地点，以及设计小组会前需完成的准备工作。",
        "owner": "Peter/Keith",
    },
]


# =================================================================
# TASK 1: Update the Meeting Minutes Tracking Record xlsx
# =================================================================
def task1_update_tracker():
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    src = "Meeting Minutes Tracking Record 0701.xlsx"
    dst = "Meeting Minutes Tracking Record 0703.xlsx"
    shutil.copy(src, dst)

    wb = load_workbook(dst)
    ws = wb["0701"]

    # Determine where to start appending
    start_row = ws.max_row + 1
    print(f"Task 1: appending from row {start_row}")

    thin = Side(style="thin", color="BFC4CC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    body_font = Font(name="Microsoft YaHei", size=10)
    align = Alignment(horizontal="left", vertical="top", wrap_text=True)

    def write_row(row_idx, date, originator, theme, progress, status, owner):
        vals = [date, originator, theme, progress, status, owner]
        for i, v in enumerate(vals, 1):
            c = ws.cell(row=row_idx, column=i, value=v)
            c.font = body_font
            c.alignment = align
            c.border = border
        ws.row_dimensions[row_idx].height = 100

    # ---- Add 3 missing Outstanding Questions ----
    for oq in MISSING_OQ:
        originator = f"SW\n{oq['section_en']}\n\n{oq['section_cn']}"
        theme = (
            f"{oq['title_en']}\n\n"
            f"{oq['title_cn']}\n\n"
            f"{oq['body_en']}\n\n"
            f"{oq['body_cn']}"
        )
        write_row(start_row, "07/01", originator, theme, None, "NO", oq["owner"])
        start_row += 1

    # ---- Add 59 Risk Register items ----
    for r in RISKS:
        rid, cat_en, cat_cn, desc_en, desc_cn, resp_en, resp_cn, owner, closed, rel_oq, _new = r
        originator = f"SW - Risk Register\n{cat_en}\n\n{cat_cn}"
        # Build theme text
        theme_parts = [f"[{rid}] {desc_en}", "", desc_cn]
        if rel_oq:
            theme_parts.append("")
            theme_parts.append(f"Related Outstanding Questions ｜ 关联待澄清问题：{rel_oq}")
        theme = "\n".join(theme_parts)
        # Build progress text (if Keith already provided a response)
        progress = None
        if resp_en:
            progress = (
                f"Keith / SW response:\n{resp_en}\n\n"
                f"Keith / SW 答复：\n{resp_cn}"
            )
        status = "YES" if closed else "NO"
        write_row(start_row, "07/02", originator, theme, progress, status, owner)
        start_row += 1

    wb.save(dst)
    print(f"Saved: {dst}")


# =================================================================
# TASK 2: Bilingual Word version of the Risk Register
# =================================================================
def task2_risk_register_word():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY   = RGBColor(0x0F, 0x34, 0x60)
    ORANGE = RGBColor(0xD9, 0x77, 0x06)
    GRAY   = RGBColor(0x47, 0x55, 0x69)
    LIGHT_GRAY = RGBColor(0x94, 0xA3, 0xB8)
    DARK   = RGBColor(0x1F, 0x29, 0x37)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN  = RGBColor(0x16, 0xA3, 0x4A)
    RED    = RGBColor(0xDC, 0x26, 0x26)

    def set_cell_bg(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    def set_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None, italic=False):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color
        r = run._element
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            r.insert(0, rpr)
        r_fonts = rpr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = OxmlElement('w:rFonts')
            rpr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), name)
        r_fonts.set(qn('w:hAnsi'), name)
        r_fonts.set(qn('w:eastAsia'), name)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), "Microsoft YaHei")
    rfonts.set(qn('w:hAnsi'), "Microsoft YaHei")
    rfonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    # ==== TITLE BANNER ====
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    cell.width = Cm(17)
    set_cell_bg(cell, "16213E")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', 240), ('bottom', 240), ('left', 300), ('right', 300)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("T-2502 Project Risk Register")
    set_font(r1, size=17, bold=True, color=WHITE)
    p.add_run("\n")
    r2 = p.add_run("Sandwich Continuous Hydrogenation Skid ｜ Bilingual Reference (EN / CN)")
    set_font(r2, size=11, color=RGBColor(0xCB, 0xD5, 0xE1))
    p.add_run("\n")
    r3 = p.add_run("项目风险登记表 · 中英双语参考版")
    set_font(r3, size=11, color=RGBColor(0xCB, 0xD5, 0xE1))

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(10)
    p_meta.paragraph_format.space_after = Pt(8)
    r_meta = p_meta.add_run(
        "📅  Source: T-2502 Project Risk Register v0.1 from Keith  ·  "
        "Translated: 2026-07-03  ·  "
        "🎯  For alignment discussion with China design team"
    )
    set_font(r_meta, size=9, color=LIGHT_GRAY, italic=True)

    # Purpose block
    tbl_pur = doc.add_table(rows=1, cols=1)
    tbl_pur.autofit = False
    tbl_pur.columns[0].width = Cm(17)
    cell_pur = tbl_pur.cell(0, 0)
    cell_pur.width = Cm(17)
    set_cell_bg(cell_pur, "F5F7FA")
    tcPr = cell_pur._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', 180), ('bottom', 180), ('left', 240), ('right', 240)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    p_pur = cell_pur.paragraphs[0]
    p_pur.paragraph_format.space_after = Pt(6)
    p_pur.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_pur.paragraph_format.line_spacing = 1.35
    r_pt = p_pur.add_run("Purpose ｜ 文件目的\n")
    set_font(r_pt, size=10.5, bold=True, color=NAVY)

    r_pe = p_pur.add_run(
        "This document is the bilingual reference version of the T-2502 Project Risk Register (v0.1) "
        "provided by Keith. It preserves all 59 risk items in their original English form and provides "
        "a Chinese translation for each. Where a response has already been recorded on Keith's side, "
        "it is retained below the risk description. The document is intended as a shared working baseline "
        "for the China design team to align on responses and residual-risk assessment."
    )
    set_font(r_pe, size=10.5, color=DARK)
    p_pur.add_run("\n\n")

    r_pc = p_pur.add_run(
        "本文件是 Keith 提供的 T-2502 项目风险登记表 (v0.1) 的中英双语参考版本。原始 59 项风险的英文表述完整保留，"
        "并对每项提供中文翻译。若 SW 侧已有答复，则保留在风险描述下方。本文件作为国内设计团队开展答复对齐"
        "与残余风险评估的共同工作基线。"
    )
    set_font(r_pc, size=10, color=GRAY)

    doc.add_paragraph()

    # Group risks by category
    from collections import OrderedDict
    grouped = OrderedDict()
    for r in RISKS:
        cat_en = r[1]
        cat_cn = r[2]
        key = (cat_en, cat_cn)
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(r)

    for (cat_en, cat_cn), items in grouped.items():
        # Section header banner
        tbl_h = doc.add_table(rows=1, cols=1)
        tbl_h.autofit = False
        tbl_h.columns[0].width = Cm(17)
        c_h = tbl_h.cell(0, 0)
        c_h.width = Cm(17)
        set_cell_bg(c_h, "0F3460")
        tcPr = c_h._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', 120), ('bottom', 120), ('left', 200), ('right', 200)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        p_h = c_h.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_he = p_h.add_run(cat_en)
        set_font(r_he, size=13, bold=True, color=WHITE)
        p_h.add_run("\n")
        r_hc = p_h.add_run(cat_cn)
        set_font(r_hc, size=11, color=RGBColor(0xCB, 0xD5, 0xE1))

        # Spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)

        for r in items:
            rid, _, _, desc_en, desc_cn, resp_en, resp_cn, owner, closed, rel_oq, _new = r

            # Title paragraph with RID
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(6)
            p_title.paragraph_format.space_after = Pt(3)
            p_title.paragraph_format.keep_with_next = True

            run_id = p_title.add_run(f"[{rid}]  ")
            set_font(run_id, size=10.5, bold=True, color=NAVY)
            # Status badge
            if closed:
                status_run = p_title.add_run("[Closed 已闭环] ")
                set_font(status_run, size=9, bold=True, color=GREEN)
            else:
                status_run = p_title.add_run("[Open 开放] ")
                set_font(status_run, size=9, bold=True, color=RED)
            if rel_oq:
                oq_run = p_title.add_run(f"[Ref: {rel_oq}] ")
                set_font(oq_run, size=9, bold=False, color=LIGHT_GRAY)

            # English description
            p_desc_en = doc.add_paragraph()
            p_desc_en.paragraph_format.space_before = Pt(0)
            p_desc_en.paragraph_format.space_after = Pt(3)
            p_desc_en.paragraph_format.left_indent = Cm(0.5)
            p_desc_en.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_desc_en.paragraph_format.line_spacing = 1.35
            rd = p_desc_en.add_run(desc_en)
            set_font(rd, size=10.5, color=DARK)

            # Chinese translation
            p_desc_cn = doc.add_paragraph()
            p_desc_cn.paragraph_format.space_before = Pt(0)
            p_desc_cn.paragraph_format.space_after = Pt(4)
            p_desc_cn.paragraph_format.left_indent = Cm(0.5)
            p_desc_cn.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_desc_cn.paragraph_format.line_spacing = 1.35
            rdc = p_desc_cn.add_run(desc_cn)
            set_font(rdc, size=10, color=GRAY)

            # Keith's response if present
            if resp_en:
                # Response header
                p_rh = doc.add_paragraph()
                p_rh.paragraph_format.space_before = Pt(2)
                p_rh.paragraph_format.space_after = Pt(2)
                p_rh.paragraph_format.left_indent = Cm(0.5)
                r_rh = p_rh.add_run("↳  Response from Keith / SW  ｜  Keith / SW 答复")
                set_font(r_rh, size=9.5, bold=True, color=ORANGE)

                # English response
                p_re = doc.add_paragraph()
                p_re.paragraph_format.space_before = Pt(0)
                p_re.paragraph_format.space_after = Pt(2)
                p_re.paragraph_format.left_indent = Cm(0.8)
                p_re.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_re.paragraph_format.line_spacing = 1.3
                rre = p_re.add_run(resp_en)
                set_font(rre, size=10, color=DARK, italic=True)

                # Chinese response
                p_rc = doc.add_paragraph()
                p_rc.paragraph_format.space_before = Pt(0)
                p_rc.paragraph_format.space_after = Pt(6)
                p_rc.paragraph_format.left_indent = Cm(0.8)
                p_rc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_rc.paragraph_format.line_spacing = 1.3
                rrc = p_rc.add_run(resp_cn)
                set_font(rrc, size=9.5, color=GRAY, italic=True)

            # Owner
            p_owner = doc.add_paragraph()
            p_owner.paragraph_format.space_before = Pt(0)
            p_owner.paragraph_format.space_after = Pt(8)
            p_owner.paragraph_format.left_indent = Cm(0.5)
            r_ow = p_owner.add_run(f"👤 Owner: {owner}")
            set_font(r_ow, size=9, color=LIGHT_GRAY, italic=True)

            # Separator
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(6)
            p_pr = sep._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'E2E8F0')
            pBdr.append(bottom)
            p_pr.append(pBdr)

    doc.save("Risk_Register_Bilingual.docx")
    print("Saved: Risk_Register_Bilingual.docx")


# =================================================================
# TASK 3: Internal Priority Warning
# =================================================================
NEW_AWARENESS_RISKS = [r for r in RISKS if r[10]]  # is_new_awareness

def task3_internal_warning():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY   = RGBColor(0x0F, 0x34, 0x60)
    ORANGE = RGBColor(0xD9, 0x77, 0x06)
    GRAY   = RGBColor(0x47, 0x55, 0x69)
    LIGHT_GRAY = RGBColor(0x94, 0xA3, 0xB8)
    DARK   = RGBColor(0x1F, 0x29, 0x37)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    RED    = RGBColor(0xDC, 0x26, 0x26)
    RED_DARK = RGBColor(0xB9, 0x1C, 0x1C)

    def set_cell_bg(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    def set_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None, italic=False):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color
        r = run._element
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            r.insert(0, rpr)
        r_fonts = rpr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = OxmlElement('w:rFonts')
            rpr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), name)
        r_fonts.set(qn('w:hAnsi'), name)
        r_fonts.set(qn('w:eastAsia'), name)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), "Microsoft YaHei")
    rfonts.set(qn('w:hAnsi'), "Microsoft YaHei")
    rfonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    # Title
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    cell.width = Cm(17)
    set_cell_bg(cell, "B91C1C")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', 240), ('bottom', 240), ('left', 300), ('right', 300)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("⚠  内部优先关注预警 — 国内团队尚未充分考虑的关键风险")
    set_font(r1, size=15, bold=True, color=WHITE)
    p.add_run("\n")
    r2 = p.add_run("Internal Priority Warning — Key Risks Not Yet Fully Considered by CN Team")
    set_font(r2, size=10.5, color=RGBColor(0xFE, 0xE2, 0xE2))
    p.add_run("\n")
    r3 = p.add_run("来源：Keith T-2502 Project Risk Register v0.1")
    set_font(r3, size=9.5, color=RGBColor(0xFE, 0xE2, 0xE2), italic=True)

    # Intro
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(12)
    p_intro.paragraph_format.space_after = Pt(10)
    p_intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_intro.paragraph_format.line_spacing = 1.4
    r_i = p_intro.add_run(
        "以下 " + str(len(NEW_AWARENESS_RISKS)) + " 条风险是在对齐 Keith 的 Risk Register 时，"
        "国内团队之前未在 Outstanding Questions 中系统覆盖、但在 UK 合规与运行体系下具有较高"
        "关注度的议题。建议由指定负责人在本周内梳理内部立场，并纳入下次双周例会的重点讨论议程。"
    )
    set_font(r_i, size=10.5, color=DARK)

    p_intro.add_run("\n\n")
    r_ie = p_intro.add_run(
        "The following " + str(len(NEW_AWARENESS_RISKS)) + " risks were surfaced when aligning "
        "against Keith's Risk Register. They were not systematically covered in the previous "
        "Outstanding Questions list, but carry high visibility under UK compliance and operating "
        "frameworks. Owners are requested to align internal positions within this week and to "
        "include these items as priority discussion points at the next biweekly meeting."
    )
    set_font(r_ie, size=9.5, color=GRAY, italic=True)

    # Priority summary table
    p_ttl = doc.add_paragraph()
    p_ttl.paragraph_format.space_before = Pt(4)
    p_ttl.paragraph_format.space_after = Pt(4)
    r_ttl = p_ttl.add_run("一、优先级速览 ｜ Priority summary")
    set_font(r_ttl, size=12, bold=True, color=NAVY)

    tbl_sum = doc.add_table(rows=len(NEW_AWARENESS_RISKS) + 1, cols=4)
    tbl_sum.autofit = False
    tbl_sum.columns[0].width = Cm(1.5)
    tbl_sum.columns[1].width = Cm(4.0)
    tbl_sum.columns[2].width = Cm(7.5)
    tbl_sum.columns[3].width = Cm(4.0)

    # Header
    hdr = tbl_sum.rows[0]
    for i, h in enumerate(["ID", "主题 Category", "简要 Brief", "负责人 Owner"]):
        c = hdr.cells[i]
        set_cell_bg(c, "0F3460")
        c.paragraphs[0].text = ""
        run = c.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=WHITE)

    for i, r in enumerate(NEW_AWARENESS_RISKS, start=1):
        rid, cat_en, cat_cn, desc_en, desc_cn, _, _, owner, _, _, _ = r
        row = tbl_sum.rows[i]
        row.cells[0].paragraphs[0].add_run(rid)
        row.cells[1].paragraphs[0].add_run(f"{cat_cn}\n{cat_en}")
        # Brief: first sentence of CN + EN
        brief_cn = desc_cn.split("。")[0][:80]
        brief_en = desc_en.split(".")[0][:80]
        row.cells[2].paragraphs[0].add_run(f"{brief_cn}\n{brief_en}")
        row.cells[3].paragraphs[0].add_run(owner)
        for cell_i in range(4):
            for run in row.cells[cell_i].paragraphs[0].runs:
                set_font(run, size=9, color=DARK)

    doc.add_paragraph()

    # Details
    p_dtl = doc.add_paragraph()
    p_dtl.paragraph_format.space_before = Pt(6)
    p_dtl.paragraph_format.space_after = Pt(4)
    r_dtl = p_dtl.add_run("二、逐条详细说明与建议行动 ｜ Detailed description and recommended actions")
    set_font(r_dtl, size=12, bold=True, color=NAVY)

    # Detailed reasoning for why each is important
    priority_hints = {
        "R2": ("为什么这是关键点：L:G 比是分离器 sizing、取样系统、传质分析的底层输入，未定义会让下游多个设计变成猜。"
               "建议行动：邵峰主导，跟研发一起在本周内明确 L:G 比范围，锁定 worst-case 值。",
               "Why critical: L:G ratio is the foundational input for separator sizing, sampler, and mass transfer. If undefined, downstream design becomes guesswork. "
               "Action: Shaofeng to lead alignment with R&D this week; lock down worst-case L:G ratio."),
        "R5": ("为什么这是关键点：UK LOLER 法规要求所有需要吊运的重物设备必须评估起重方式，无 lifting aid 会阻碍设备维护。"
               "建议行动：赵子亮 / 范双双 在 3D 模型中标注需要吊臂的位置，加入 davit / lifting beam 设计。",
               "Why critical: UK LOLER regulation requires lifting evaluation for all heavy equipment; absence of lifting aids will prevent maintenance. "
               "Action: Zhao/Fan to mark lifting aid positions on 3D model and add davit / lifting beam design."),
        "R6": ("为什么这是关键点：现场照明设计不足会影响维护和操作安全，且是 UK 合规检查点。"
               "建议行动：CFCT 在总布置中加入照明布置图，明确操作/维护区照度水平。",
               "Why critical: Inadequate lighting affects maintenance and operational safety; a UK compliance checkpoint. "
               "Action: CFCT to add lighting layout to general arrangement drawing with illumination levels defined."),
        "R7": ("为什么这是关键点：氢气分子小，水压试验证明不了氢密封性。UK 通常要求氦气检漏，国内加工厂不常做。这可能是重大成本 / 工期问题。"
               "建议行动：赵子亮尽快与加工厂确认氦检漏能力和成本，若国内无法做则可能要在 UK 现场做，需要提前规划。",
               "Why critical: H₂ molecules are small; hydraulic testing cannot demonstrate H₂ integrity. UK typically requires helium leak testing which China fabricators rarely perform. Potentially major cost / schedule risk. "
               "Action: Zhao to confirm helium test capability and cost with fabricator immediately; if not available in China, plan for UK-site testing."),
        "R8": ("为什么这是关键点：紧急泄放总管尺寸如果算错会导致背压 > 10% set，PSV 型号必须换成 balanced-bellows 或 pilot 操作型，牵动整个 PSV 采购。"
               "建议行动：胡超群按 API 521 §7 核算 vent header 尺寸，并把结果同步 Keith 那边的 header 设计。",
               "Why critical: Incorrectly sized vent header leads to backpressure >10% set → forces balanced-bellows or pilot-operated PSVs, disrupting entire PSV procurement. "
               "Action: Hu to size vent header per API 521 §7 and synchronise with Keith's header design."),
        "R9": ("为什么这是关键点：保温需求需要预留空间和支撑；后期加保温会遇到空间不够或管路 heat loss 过大。"
               "建议行动：范双双确认所有工艺管路和设备的保温需求（是 / 否 / 厚度），纳入 3D 模型审查。",
               "Why critical: Insulation needs both space and support; retrofitting insulation later hits space constraints or excessive heat loss. "
               "Action: Fan to confirm insulation requirements (yes/no/thickness) for all process pipes and equipment; include in 3D model review."),
        "R11": ("为什么这是关键点：如果撬块在 UK 现场组装，中方 CE 责任人（Responsible Person）可能无法为整撬发 Declaration of Conformity。这是 CE 合规的根本性问题。"
                "建议行动：赵子亮与 SGS（认证方）明确 —— 撬块必须在中国 100% 完成组装再运，还是可以现场组装但由 UK 一方担任 CE 责任人。这个决定倒逼加工工艺。",
                "Why critical: If skid is assembled on-site in UK, the China-based CE Responsible Person may not be able to issue Declaration of Conformity. Fundamental CE compliance issue. "
                "Action: Zhao to clarify with SGS — must skid be 100% assembled in China before shipping, or can site assembly occur with UK-side as CE Responsible Person? This decision drives fabrication approach."),
        "R16": ("为什么这是关键点：密闭加料系统防止暴露 OEB 3 级催化剂粉尘；无设计的话，未来加料操作会因职业卫生 (COSHH) 问题被卡住。"
                "建议行动：范双双或 CIMT 提供密闭加料系统的原理图和空间预留方案。",
                "Why critical: Contained charging system prevents exposure to OEB 3 catalyst dust; without design, future charging operations will be blocked by occupational hygiene (COSHH) issues. "
                "Action: Fan or CIMT to provide contained charging system schematic and space allocation."),
        "R24": ("为什么这是关键点：SE01 脱气不完全会让 H₂ 进入 BT04 常压罐 → 常压罐里累积可燃气 = 潜在爆炸风险。这是 SE01 设计的一个安全隐患。"
                "建议行动：胡超群 / 范双双评估 BT04 顶部是否需要独立的 vent 或 N₂ blanketing 加强，以及 SE01 液相出口的气相夹带量估算。",
                "Why critical: Incomplete degassing in SE01 lets H₂ into atmospheric BT04 → flammable gas accumulation → potential explosion hazard. Safety concern for SE01 design. "
                "Action: Hu/Fan to evaluate whether BT04 needs independent vent or enhanced N₂ blanketing; also estimate carry-over gas at SE01 liquid outlet."),
        "R29": ("为什么这是关键点：取样瓶如果气相入口速率过高，可能超压。需要在设计上加限流或防超压设施。"
                "建议行动：胡超群 / 高宇 明确取样系统防超压装置（限流孔 / 安全阀 / 压力控制）。",
                "Why critical: Sample bottle may over-pressurise if gas inlet rate is high; requires flow restrictor or over-pressure protection. "
                "Action: Hu/Yu to define over-pressure protection for the sample system (restriction orifice / PSV / pressure control)."),
        "R31": ("为什么这是关键点：密闭卸料系统需要顶部空间，若设计不留则催化剂更换必须敞开进行，违反职业卫生要求。"
                "建议行动：范双双在总布置中确认密闭卸料所需顶部净空。",
                "Why critical: Contained discharging system requires overhead clearance; if not designed, catalyst replacement must be done open, violating occupational hygiene. "
                "Action: Fan to confirm required overhead clearance for contained discharge in general arrangement."),
        "R34": ("为什么这是关键点：法兰缝隙残留是 batch-to-batch 交叉污染的经典机制。清洗在常压下做，法兰间隙微开时藏物，高压运行时被挤出污染下一批。"
                "建议行动：胡超群 / 范双双评估是否需要在关键法兰位置增加拆装清洗要求或使用不同密封结构。",
                "Why critical: Flange gap contamination is a classic batch-to-batch cross-contamination mechanism. Cleaning done at atmospheric while flanges are slightly separated hides material; high-pressure operation then extrudes it into next batch. "
                "Action: Hu/Fan to evaluate whether critical flanges need disassembly cleaning requirements or different seal design."),
        "R36": ("为什么这是关键点：设备布置未评估旁路 / 死管 / 排放 / 放空 —— 这些是清洗死角常见来源。"
                "建议行动：范双双在 3D 模型审查会上专门标记所有 dead legs、bypass 管段，并逐个评估清洗方式。",
                "Why critical: Equipment layout has not been evaluated for bypasses / dead legs / drains / vents — common sources of cleaning dead zones. "
                "Action: Fan to specifically mark all dead legs and bypass segments in 3D model review, evaluating cleaning method for each."),
        "R58": ("为什么这是关键点：UK EA (Environment Agency) 许可要求 A/B 级 VOC 年排放量证明。目前无。若未在 CE Technical File 中体现，可能会推迟运行许可。"
                "建议行动：邵峰 / 胡超群 用工艺物料衡算估算年 VOC 排放量，纳入 Technical File。",
                "Why critical: UK EA (Environment Agency) permit requires annual VOC (Class A/B) emission proof. Not currently available. If not shown in CE Technical File, could delay operating permit. "
                "Action: Shaofeng/Hu to estimate annual VOC emission from process material balance and include in Technical File."),
        "R59": ("为什么这是关键点：Keith 一次性列了 10 项英国合规要求（CE / Functional Safety / DSEAR / ATEX / PSSR / COSHH / PUWER / LOLER / Lighting / Occupational Hygiene / EA Permit）。"
                "国内团队需要有一份合规 checklist 来系统覆盖。建议行动：Peter 起草项目 UK 合规 checklist，每一项落实到具体的责任人和交付物。",
                "Why critical: Keith listed 10 UK compliance requirements in one item. CN team needs a compliance checklist to systematically cover them. "
                "Action: Peter to draft UK compliance checklist for the project, mapping each item to a specific owner and deliverable."),
    }

    for r in NEW_AWARENESS_RISKS:
        rid, cat_en, cat_cn, desc_en, desc_cn, resp_en, resp_cn, owner, _, _, _ = r

        # Item header
        p_ih = doc.add_paragraph()
        p_ih.paragraph_format.space_before = Pt(10)
        p_ih.paragraph_format.space_after = Pt(3)
        p_ih.paragraph_format.keep_with_next = True
        r_ih_id = p_ih.add_run(f"[{rid}]  ")
        set_font(r_ih_id, size=11, bold=True, color=RED_DARK)
        r_ih_cat = p_ih.add_run(f"{cat_cn}  ｜  {cat_en}")
        set_font(r_ih_cat, size=11, bold=True, color=NAVY)

        # Risk description CN then EN
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(3)
        p_d.paragraph_format.left_indent = Cm(0.5)
        p_d.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_d.paragraph_format.line_spacing = 1.35
        r_dt = p_d.add_run("风险描述：")
        set_font(r_dt, size=9.5, bold=True, color=ORANGE)
        p_d.add_run("\n")
        rd = p_d.add_run(desc_cn)
        set_font(rd, size=10, color=DARK)
        p_d.add_run("\n\n")
        rde = p_d.add_run(desc_en)
        set_font(rde, size=9.5, color=GRAY, italic=True)

        # Hint (recommended action)
        if rid in priority_hints:
            hint_cn, hint_en = priority_hints[rid]
            p_h = doc.add_paragraph()
            p_h.paragraph_format.space_before = Pt(2)
            p_h.paragraph_format.space_after = Pt(2)
            p_h.paragraph_format.left_indent = Cm(0.5)
            p_h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_h.paragraph_format.line_spacing = 1.35
            r_hh = p_h.add_run("💡 分析与建议行动：")
            set_font(r_hh, size=9.5, bold=True, color=ORANGE)
            p_h.add_run("\n")
            rh_cn = p_h.add_run(hint_cn)
            set_font(rh_cn, size=10, color=DARK)
            p_h.add_run("\n\n")
            rh_en = p_h.add_run(hint_en)
            set_font(rh_en, size=9.5, color=GRAY, italic=True)

        # Owner
        p_o = doc.add_paragraph()
        p_o.paragraph_format.space_before = Pt(2)
        p_o.paragraph_format.space_after = Pt(8)
        p_o.paragraph_format.left_indent = Cm(0.5)
        r_o = p_o.add_run(f"👤 建议负责人 Suggested Owner: {owner}")
        set_font(r_o, size=9.5, color=LIGHT_GRAY, italic=True)

        # Separator
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(4)
        p_pr = sep._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'FCA5A5')
        pBdr.append(bottom)
        p_pr.append(pBdr)

    # Footer with recommended actions
    doc.add_paragraph()
    tbl_end = doc.add_table(rows=1, cols=1)
    tbl_end.autofit = False
    tbl_end.columns[0].width = Cm(17)
    ec = tbl_end.cell(0, 0)
    ec.width = Cm(17)
    set_cell_bg(ec, "FEE2E2")
    tcPr = ec._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', 180), ('bottom', 180), ('left', 240), ('right', 240)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p_e = ec.paragraphs[0]
    p_e.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_e.paragraph_format.line_spacing = 1.4
    r_et = p_e.add_run("📋 后续动作 ｜ Next steps\n\n")
    set_font(r_et, size=11, bold=True, color=RED_DARK)

    r_e_zh = p_e.add_run(
        "1. 本周内由指定负责人梳理内部立场，形成初步答复\n"
        "2. 下次双周例会（7 月中旬）把这 " + str(len(NEW_AWARENESS_RISKS)) +
        " 条设为优先讨论议题\n"
        "3. 讨论完成后在 Meeting Minutes Tracking Record 中记录 China Response 与 Residual Risk\n"
        "4. R11 (CE 责任人归属) 和 R7 (氦检漏) 建议单独启动决策，涉及成本与工期影响较大\n"
    )
    set_font(r_e_zh, size=10, color=RGBColor(0x7F, 0x1D, 0x1D))

    r_e_en = p_e.add_run(
        "\n"
        "1. Owners to establish internal positions this week and prepare initial responses\n"
        "2. Set these " + str(len(NEW_AWARENESS_RISKS)) + " items as priority topics at the next biweekly meeting (mid-July)\n"
        "3. After discussion, record China Response and Residual Risk in Meeting Minutes Tracking Record\n"
        "4. R11 (CE Responsible Person) and R7 (helium leak testing) may need separate decisions due to significant cost / schedule impact"
    )
    set_font(r_e_en, size=9.5, color=RGBColor(0x7F, 0x1D, 0x1D), italic=True)

    doc.save("Internal_Priority_Warning_Risks.docx")
    print("Saved: Internal_Priority_Warning_Risks.docx")


# =================================================================
# Run all three
# =================================================================
os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
task1_update_tracker()
task2_risk_register_word()
task3_internal_warning()

# Report the priority risks flagged
print(f"\n=== NEW_AWARENESS priority risks ({len(NEW_AWARENESS_RISKS)}) ===")
for r in NEW_AWARENESS_RISKS:
    print(f"  {r[0]}: {r[2]} - {r[4][:60]}")
