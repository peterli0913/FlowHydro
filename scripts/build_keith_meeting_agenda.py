"""
Build a bilingual one-page agenda (landscape A4) for the upcoming Keith
technical review, ordered by the row order of
"Meeting Minutes Tracking Record 0706.xlsx" (sheet 0701).

Only the items that still need to be discussed / clarified / decided with
Keith are included (closed "Keith/SW response" items are omitted).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Microsoft YaHei"
C_TITLE = RGBColor(0x00, 0x36, 0x6A)
C_BAND = "1F4E79"      # section band fill
C_BAND_TXT = RGBColor(0xFF, 0xFF, 0xFF)
C_HEAD = "00366A"
C_TAG = {
    "澄": RGBColor(0xC0, 0x50, 0x00),   # clarify - orange
    "决": RGBColor(0xB0, 0x00, 0x00),   # decide  - red
    "复": RGBColor(0x1F, 0x4E, 0x79),   # respond - blue
    "报": RGBColor(0x55, 0x6070 // 256, 0x6B),  # inform - gray-ish (fix below)
}
C_TAG = {
    "澄": RGBColor(0xC0, 0x50, 0x00),
    "决": RGBColor(0xB0, 0x00, 0x00),
    "复": RGBColor(0x1F, 0x4E, 0x79),
    "报": RGBColor(0x70, 0x70, 0x70),
}
TAG_LABEL = {
    "澄": "澄清 Clarify",
    "决": "决策 Decide",
    "复": "回复 Respond",
    "报": "通报 Inform",
}

# ---- agenda content, in ROW ORDER of the tracking sheet --------------------
# (section, [ (row, tag, topic_cn/en, point_cn, point_en, owner) ... ])
SECTIONS = [
    ("3D 布置与空间  ·  3D Layout & Space", [
        ("R31", "决", "放空缓冲罐 T06 布置 / Vent buffer tank T06",
         "T06 移至 service area；撬块能否加宽 150–200 mm？",
         "Move T06 to service area; can skid width +150–200 mm?", "Ziliang"),
        ("R44", "决", "操作通道净空 / Operator access clearance",
         "当前仅 300–350 mm；请确认最小净空标准",
         "Only 300–350 mm now; confirm min clearance target", "Fan/All"),
        ("R46/47", "澄", "管嘴方位 & 催化剂加卸料 / Nozzles & catalyst charge",
         "空间受限维护不便；加卸料口在端面，确认空间是否足够",
         "Space-limited; charge/discharge ports on end face — access ok?",
         "Fan/All"),
        ("R48", "决", "撬块–公用工程接口 / Tie-in points",
         "缺仪表空气/电气/电缆穿墙接口，需对齐位置与规格",
         "Missing IA / power / cable interfaces; align location & spec",
         "Fan/All"),
        ("R49", "澄", "穿墙电缆孔 / Cable wall penetrations",
         "14 孔是否够、孔径与最大电缆兼容性",
         "Confirm 14-slot count, bore size & cable compatibility",
         "Dezhi/Shaofeng"),
    ]),
    ("控制策略  ·  Control Philosophy", [
        ("R50/51", "报", "控制策略 & 联锁清单 / Control Philosophy & interlocks",
         "初版完成、校核中，预计下周发出（含硬件 vs 软件联锁）",
         "Drafts done, under review, issue next week (incl. HW/SW interlocks)",
         "Chaoqun/Dezhi"),
        ("R52", "澄", "操作模式与故障位 / Modes & fail positions",
         "请 Keith 定义各模式、故障位、故障→正常恢复方式",
         "Keith to define modes, fail positions & recovery", "Chaoqun"),
        ("R53", "澄", "模式切换保护 / Mode-transition safeguards",
         "听 Keith 需求；R52 讨论后 1 周反馈",
         "Hear needs; feedback 1 wk after R52", "Dezhi"),
    ]),
    ("取样系统 / 仪表  ·  Sampler & Instrumentation", [
        ("R56", "复", "气相主导流态取样 / Gas-rich sampling regime",
         "回复取样流程：先排气卸压→N₂ 压液至取样瓶",
         "Explain procedure: vent & depressurise → N₂ push liquid",
         "Chaoqun"),
        ("R57", "澄", "导波液位计 / 超临界 / Guided-wave LT & supercritical",
         "超临界下是否仍有液相？无液相则液位计无法识别",
         "Any liquid phase under supercritical? If none, LT fails",
         "Yu/Chaoqun"),
        ("R58/59", "报", "液位开关选型 & 仪表数据单 / Level switch & datasheet",
         "仅 VEGA VEGASWING 63 满足 DN25 Cl900；数据单修订版已完成待意见",
         "Only VEGA meets DN25 Cl900; datasheet revised, awaiting comments",
         "Chaoqun/Yu"),
        ("R113", "复", "控制器差压 / PIC030 减压器 / DP & PIC030 reducer",
         "FIC004/PIC028/PIC030 差压问题；PIC030 选型与 SIL 独立性",
         "DP issues; confirm PIC030 selection & SIL independence",
         "Chaoqun/Yu"),
    ]),
    ("工程设计 / 工艺  ·  Engineering & Process", [
        ("R62", "澄", "液气比 L:G ratio",
         "请明确工况优先级，否则物料平衡只能给范围值",
         "Define priority scenarios, else mass balance stays ranges",
         "Feng/Chaoqun"),
        ("R63", "决", "催化剂规格 / OEB3 / Catalyst spec & OEB3",
         "OEB3 相对 URS 存在偏离，请确认是否接受",
         "OEB3 is a URS deviation — confirm acceptance", "Feng/Chaoqun"),
        ("R120", "决", "CR01 管径 DN80→DN50 → 产能 / capacity",
         "产能可能 < URS 30 kg/day：改 URS 还是保持 DN80？（新）",
         "May fall below URS 30 kg/day: revise URS or keep DN80? (NEW)",
         "Feng/Chaoqun"),
    ]),
    ("反应器建造  ·  Continuous Reactor Build", [
        ("R67", "决", "系统检漏 / 氦检漏 / Leak (helium) test",
         "建议现场组装后做；现场是否具备试验条件？",
         "Do on site after reassembly; site test capability?", "Ziliang"),
        ("R69", "澄", "保温 / Insulation",
         "请提供保温材料，厚度由我方计算",
         "Provide insulation material; we size the thickness",
         "Cheng/Chaoqun"),
        ("R70", "澄", "模块化设计 / Modular design",
         "请举例说明担心的现场施工工作量",
         "Clarify the site-construction workload concern", "TBD"),
        ("R71", "澄", "现场组装 vs CE/UKCA 符合性声明",
         "请澄清担心点：授权代表资质范围 or 签发流程？",
         "Clarify: authorised-rep scope or DoC issuance?", "Ziliang"),
        ("R80", "复", "精滤器堵塞 FTR9832 / Polishing filter",
         "差压表已升级 + 溶剂清洗覆盖粉尘；向厂家二次确认",
         "DP gauge upgraded + solvent clean; re-confirm w/ vendor",
         "Chaoqun"),
    ]),
    ("分离器 / 清洗  ·  Separator & Cleaning", [
        ("R84", "复", "产品脱气 / H₂ 带出 / Degassing & H₂ carryover",
         "分离器满足气液分离要求（含超临界工况）",
         "Separator meets G/L separation (incl. supercritical)",
         "Chaoqun"),
        ("R95", "澄", "清洗验证 / 确认 / Cleaning validation",
         "请 Keith 给出验收标准与验证方法期望",
         "Keith to give acceptance criteria & verification method",
         "Chaoqun"),
    ]),
    ("自动化 / 文档 / 法规  ·  Automation · Docs · Regulatory", [
        ("R115", "报", "自动化策略文档 / Automation philosophy",
         "国内专题沟通后反馈分工与文档",
         "Feedback after internal session", "Chaoqun/Dezhi/Ziliang"),
        ("R60/116", "决", "交付文件版本时间表 / Deliverable rev dates",
         "给出各文件现版本 + 下一版发出日期（HAZOP 排期依据）",
         "Provide current + next rev dates (basis for HAZOP schedule)",
         "Shaofeng/All"),
        ("R118", "澄", "排放量 / EA 许可 / Emissions & EA permit",
         "边界：我方提供到汇集管，之后由 SW 汇总",
         "Boundary: we provide up to header; SW consolidates after",
         "Chaoqun"),
        ("R119", "澄", "CE / 法规合规范围 / Regulatory scope",
         "明确 scope 分界：撬块本体 vs SW 场地事项",
         "Clarify scope split: skid package vs SW site items",
         "Ziliang/All"),
    ]),
]


# ---------- helpers ---------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_run(run, size=8, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)

def no_space(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

def set_cell_margins(cell, top=0, bottom=0, left=30, right=30):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom),
                     ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


# ---------- build -----------------------------------------------------------

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
sec.left_margin = sec.right_margin = Cm(0.7)
sec.top_margin = sec.bottom_margin = Cm(0.45)

# Title
t = doc.add_paragraph()
no_space(t)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("与 Keith 技术讨论 · 会议议程  |  Technical Review with Keith — Agenda")
set_run(r, size=11.5, bold=True, color=C_TITLE)


# Table
n_rows = 1 + sum(1 + len(items) for _, items in SECTIONS)
table = doc.add_table(rows=0, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
# column widths
widths = [Cm(1.4), Cm(1.7), Cm(8.6), Cm(13.2), Cm(2.6)]

def add_row(cells_data, *, header=False, band=False, tag=None):
    row = table.add_row()
    if header:
        row.height = Cm(0.50)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    else:
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        no_space(p)
        text = cells_data[i]
        if band:
            set_cell_bg(cell, C_BAND)
            r = p.add_run(text)
            set_run(r, size=9, bold=True, color=C_BAND_TXT)
        elif header:
            set_cell_bg(cell, C_HEAD)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run(r, size=8.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            if i == 1 and tag:  # tag cell colored
                r = p.add_run(text)
                set_run(r, size=7, bold=True, color=C_TAG.get(tag))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 0:
                r = p.add_run(text)
                set_run(r, size=7, bold=True, color=RGBColor(0x00,0x36,0x6A))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 4:
                r = p.add_run(text)
                set_run(r, size=6)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                r = p.add_run(text)
                set_run(r, size=6)
    return row

# header row
add_row(["Row", "类型\nType", "议题  Topic", "讨论要点  Discussion point",
         "负责\nOwner"], header=True)

for section_name, items in SECTIONS:
    # band row spanning all columns
    brow = table.add_row()
    brow.height = Cm(0.36)
    brow.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, cell in enumerate(brow.cells):
        cell.width = widths[i]
    # merge
    merged = brow.cells[0].merge(brow.cells[4])
    set_cell_bg(merged, C_BAND)
    set_cell_margins(merged)
    p = merged.paragraphs[0]
    no_space(p)
    r = p.add_run(section_name)
    set_run(r, size=7.5, bold=True, color=C_BAND_TXT)
    for row_no, tag, topic, pt_cn, pt_en, owner in items:
        add_row([row_no, TAG_LABEL[tag].split()[0],
                 topic, f"{pt_cn}\n{pt_en}", owner], tag=tag)

# footer note
f = doc.add_paragraph()
no_space(f)
r = f.add_run("类型 Type：澄清 Clarify · 决策 Decide · 回复 Respond · 通报 Inform     ")
set_run(r, size=6.5, italic=True, color=RGBColor(0x55, 0x55, 0x55))
r = f.add_run("｜  重点 Key focus：R120（产能 vs URS）· R62（工况优先级）· "
              "R57/R84（超临界是否有液相）· R31/R44（撬块加宽）· "
              "R60/116（交付文件版本时间表）")
set_run(r, size=6.5, bold=True, color=RGBColor(0xC0, 0x50, 0x00))

doc.save("Keith_Meeting_Agenda_0714.docx")
print("Saved: Keith_Meeting_Agenda_0714.docx")
