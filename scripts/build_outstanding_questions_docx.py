"""
Build a Word (.docx) version of the "Consolidated Outstanding Design
Questions" document that mirrors the HTML content.

Same content as outstanding_questions_for_keith.html, formatted for
professional Word export (suitable for email attachment / print).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# ============ Colors ============
NAVY   = RGBColor(0x0F, 0x34, 0x60)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
GRAY   = RGBColor(0x47, 0x55, 0x69)
LIGHT_GRAY = RGBColor(0x94, 0xA3, 0xB8)
DARK   = RGBColor(0x1F, 0x29, 0x37)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ============ Helpers ============
def set_cell_bg(cell, color_hex):
    """Set the background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Also set east-Asian font (for proper Chinese rendering in Word)
    r = run._element
    rpr = r.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r.insert(0, rpr)
    r_fonts = rpr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        rpr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), name)
    r_fonts.set(qn('w:hAnsi'), name)
    r_fonts.set(qn('w:eastAsia'), name)


def add_para(doc, text, size=10.5, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, italic=False,
             indent_left=None):
    p = doc.add_paragraph()
    p.alignment = align
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_section_header(doc, num_and_title_en, title_cn):
    """A colored section header banner using a 1-row table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    cell.width = Cm(17)
    set_cell_bg(cell, "0F3460")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Left cell padding via margins
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', 120), ('bottom', 120), ('left', 200), ('right', 200)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(num_and_title_en)
    set_font(run, size=13, bold=True, color=WHITE)
    p.add_run("\n")
    run2 = p.add_run(title_cn)
    set_font(run2, size=11, bold=False, color=RGBColor(0xCB, 0xD5, 0xE1))

    # Space after header
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.space_before = Pt(0)


def add_question(doc, qid, title_en, title_cn, body_en, body_cn):
    """Add a single question with QID pill, EN title, CN title, EN body, CN body."""
    # Title paragraph — QID pill + English + Chinese
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(3)
    p_title.paragraph_format.keep_with_next = True

    # QID pill: bold in navy background... use text simulation
    run_id = p_title.add_run(f" {qid} ")
    set_font(run_id, size=9, bold=True, color=WHITE)
    # We can't easily add fill to a run inline in python-docx.
    # Use a bracketed style instead:
    # Actually let's use "[ Q1.1 ]" format visually
    p_title.runs[0].text = f"[{qid}]  "
    set_font(p_title.runs[0], size=10, bold=True, color=NAVY)

    run_en = p_title.add_run(title_en)
    set_font(run_en, size=10.5, bold=True, color=DARK)

    p_title_cn = doc.add_paragraph()
    p_title_cn.paragraph_format.space_before = Pt(0)
    p_title_cn.paragraph_format.space_after = Pt(4)
    p_title_cn.paragraph_format.left_indent = Cm(0.6)
    p_title_cn.paragraph_format.keep_with_next = True
    run_cn = p_title_cn.add_run(title_cn)
    set_font(run_cn, size=10, bold=False, color=GRAY)

    # Body — English
    p_body_en = doc.add_paragraph()
    p_body_en.paragraph_format.space_before = Pt(0)
    p_body_en.paragraph_format.space_after = Pt(3)
    p_body_en.paragraph_format.left_indent = Cm(0.6)
    p_body_en.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_body_en.paragraph_format.line_spacing = 1.35
    run_be = p_body_en.add_run(body_en)
    set_font(run_be, size=10.5, color=DARK)

    # Body — Chinese
    p_body_cn = doc.add_paragraph()
    p_body_cn.paragraph_format.space_before = Pt(0)
    p_body_cn.paragraph_format.space_after = Pt(10)
    p_body_cn.paragraph_format.left_indent = Cm(0.6)
    p_body_cn.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_body_cn.paragraph_format.line_spacing = 1.35
    run_bc = p_body_cn.add_run(body_cn)
    set_font(run_bc, size=10, color=RGBColor(0x4B, 0x55, 0x63))

    # Thin separator under each question
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(6)
    p_pr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 0.5pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    p_pr.append(pBdr)


# ============ Build document ============
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Set default style to Microsoft YaHei
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:ascii'), "Microsoft YaHei")
rfonts.set(qn('w:hAnsi'), "Microsoft YaHei")
rfonts.set(qn('w:eastAsia'), "Microsoft YaHei")

# ==== TITLE BANNER ====
tbl = doc.add_table(rows=1, cols=1)
tbl.autofit = False
tbl.columns[0].width = Cm(17)
cell = tbl.cell(0, 0)
cell.width = Cm(17)
set_cell_bg(cell, "16213E")
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
tcPr = cell._tc.get_or_add_tcPr()
tcMar = OxmlElement('w:tcMar')
for side, val in [('top', 240), ('bottom', 240), ('left', 300), ('right', 300)]:
    node = OxmlElement(f'w:{side}')
    node.set(qn('w:w'), str(val))
    node.set(qn('w:type'), 'dxa')
    tcMar.append(node)
tcPr.append(tcMar)

p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(0)
r1 = p.add_run("Consolidated Outstanding Design Questions")
set_font(r1, size=17, bold=True, color=WHITE)
p.add_run("\n")
r2 = p.add_run("Sandwich Continuous Hydrogenation Skid — Pre-HAZOP Design Review")
set_font(r2, size=11, color=RGBColor(0xCB, 0xD5, 0xE1))
p.add_run("\n")
r3 = p.add_run("连续氢化撬块 · HAZOP 前设计评审 · 待澄清问题清单")
set_font(r3, size=11, color=RGBColor(0xCB, 0xD5, 0xE1))

# Meta info line
p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_before = Pt(10)
p_meta.paragraph_format.space_after = Pt(8)
r_meta = p_meta.add_run(
    "📅  Prepared: 2026-07-02      👥  From: IEPE / CFCT Design Team → Keith (Sandwich)      "
    "🎯  Target Alignment: prior to HAZOP session, planned mid-August 2026"
)
set_font(r_meta, size=9, color=LIGHT_GRAY, italic=True)

# ==== PURPOSE BLOCK ====
tbl_pur = doc.add_table(rows=1, cols=1)
tbl_pur.autofit = False
tbl_pur.columns[0].width = Cm(17)
cell_pur = tbl_pur.cell(0, 0)
cell_pur.width = Cm(17)
set_cell_bg(cell_pur, "F5F7FA")
tcPr = cell_pur._tc.get_or_add_tcPr()
tcMar = OxmlElement('w:tcMar')
for side, val in [('top', 180), ('bottom', 180), ('left', 240), ('right', 240)]:
    node = OxmlElement(f'w:{side}')
    node.set(qn('w:w'), str(val))
    node.set(qn('w:type'), 'dxa')
    tcMar.append(node)
tcPr.append(tcMar)

p_pur = cell_pur.paragraphs[0]
p_pur.paragraph_format.space_after = Pt(6)
p_pur.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p_pur.paragraph_format.line_spacing = 1.35
r_pt = p_pur.add_run("Purpose of this document ｜ 文件目的\n")
set_font(r_pt, size=10.5, bold=True, color=NAVY)

r_pe = p_pur.add_run(
    "Ahead of the planned HAZOP session, this document consolidates the "
    "outstanding design questions across cleaning philosophy, layout & "
    "space, control philosophy, mechanical design, sampler system, "
    "instrumentation, and documentation status. Alignment on each of "
    "the items below is requested to allow the HAZOP session to proceed "
    "on a stable design basis."
)
set_font(r_pe, size=10.5, color=DARK)

p_pur.add_run("\n\n")

r_pc = p_pur.add_run(
    "为便于计划中的 HAZOP 分析顺利进行，本文件汇总了当前在清洗策略、布置与空间、"
    "控制策略、机械设计、取样系统、仪表配置、文档状态等方面尚未澄清的设计问题。"
    "烦请在 HAZOP 会议前就以下各项达成一致，以确保 HAZOP 在稳定的设计基线上开展。"
)
set_font(r_pc, size=10, color=GRAY)

doc.add_paragraph()  # spacer

# ==============================================================
# CONTENT SECTIONS
# ==============================================================
sections = [
    {
        "num_en": "1. Cleaning Philosophy & Cleaning Sequence",
        "num_cn": "清洗策略与清洗程序",
        "questions": [
            (
                "Q1.1",
                "Overall cleaning philosophy per vessel",
                "各容器整体清洗策略",
                "Please confirm the intended cleaning philosophy for each vessel in the skid — spray-only, soak-only, spray-and-soak combined, or an alternative approach. Please also indicate whether a single philosophy applies across all vessels or whether the strategy varies by vessel.",
                "请确认撬块内各台容器的清洗策略：单独喷淋、单独浸泡、喷淋+浸泡组合，或其它方式。请同时说明是否所有容器采用同一策略，或按容器区别设计。",
            ),
            (
                "Q1.2",
                "Cleanliness verification method & acceptance criteria",
                "清洁度验证方法与验收标准",
                "Please advise the verification method(s) used to confirm cleaning completion (visual inspection, swab / rinse sample analysis, mass balance, calculated residue, etc.) and the numerical acceptance criteria applied (e.g. residue mass, ppm concentration threshold).",
                "请说明用于确认清洗完成的验证方法（目视检查 / 拭子或冲洗样分析 / 物料衡算 / 残余物计算等），以及所采用的数值验收标准（如残余物质量、浓度阈值 ppm）。",
            ),
            (
                "Q1.3",
                "Sight glass availability for SE01",
                "SE01 视镜配置",
                "The reactor vessels are understood to be fitted with sight glasses to support visual verification of cleanliness. The gas–liquid separator SE01, however, does not include a sight glass. Please confirm the intended method for verifying cleanliness on SE01, and whether the addition of a sight glass is being considered as a design change.",
                "反应柱已配置视镜以支持目视清洁度检查，但气液分离器 SE01 未配置视镜。请确认 SE01 清洁度验证的具体手段，以及是否考虑将新增视镜作为设计变更纳入。",
            ),
            (
                "Q1.4",
                "Spray ball flow rate vs. vessel fill time",
                "喷淋球流量与容器充液时间的匹配",
                "The lowest available spray ball flow rate (approximately 2 L/min) results in SE01 reaching the high-level trip point in under 25 seconds. This limits the time window in which mechanical spray action can achieve wetting coverage. Please advise the intended cleaning-time window on SE01 and the approach for ensuring adequate mechanical cleaning action within the design constraint.",
                "现有最小可选喷淋球流量约 2 L/min，导致 SE01 容器在 25 秒内即达到高液位跳停，机械喷淋作用的有效时间窗口非常短。请说明 SE01 上预期的清洗时间窗口，以及在此设计约束下如何确保获得充分的机械清洗作用。",
            ),
            (
                "Q1.5",
                "Drain line sizing and fill–drain cycle",
                "排放管径与充液-排空循环",
                "Under the current drain line configuration, the drain rate is significantly slower than the spray inflow rate. Please advise the intended drain line diameter (whether upsizing is planned), the estimated single-cycle fill and drain durations, and the number of cycles anticipated per cleaning campaign.",
                "当前排放管径下，排放速率显著慢于喷淋进料速率。请说明拟采用的排放管径（是否放大）、单次循环的充液和排空时间估算，以及每次清洗预期需要的循环次数。",
            ),
            (
                "Q1.6",
                "Soak cleaning vs. high-level interlock defeat mechanism",
                "浸泡清洗与高液位联锁绕过机制",
                "If soak cleaning is adopted on SE01, the liquid level during cleaning will exceed the high-level safety interlock (located at approximately one-third of vessel height). Please specify: (a) the physical mechanism intended for defeating this interlock during cleaning (key switch, procedural lockout, or other); (b) the authorisation and access controls surrounding the defeat mechanism; and (c) the safeguards against inadvertent defeat during normal processing operation.",
                "若 SE01 采用浸泡清洗，清洗过程中的液位将超过位于容器约 1/3 高度的高液位安全联锁。请说明：(a) 清洗时用于绕过（defeat）该联锁的物理机制（钥匙开关 / 程序锁定 / 其它）；(b) 该绕过机制的授权与权限控制方式；(c) 防止正常运行时意外绕过的保护措施。",
            ),
            (
                "Q1.7",
                "Cleaning of the guided-wave level transmitter housing",
                "Guided wave 液位变送器外壳的清洁",
                "The guided-wave level transmitter housing (approximately 100 mm diameter, side-mounted on SE01) sits above the high-level switch and lies outside the spray ball coverage envelope. Please advise the intended cleaning approach for this housing, and any mechanical or procedural provisions required to ensure it can be cleaned to the acceptance criteria in Q1.2.",
                "液位变送器外壳（约 100 mm 直径，SE01 侧接管上）位于高液位开关之上，不在喷淋球覆盖范围内。请说明该外壳的清洁方案，以及为达到 Q1.2 中的验收标准所需的机械或程序性措施。",
            ),
            (
                "Q1.8",
                "Reactor cleaning cycle time and manual interventions",
                "反应器清洗循环时间与人工干预",
                "For the 200-L reactor vessels, a single spray-and-drain cycle involves an extended drain phase relative to the spray phase. Multiple cycles are expected to achieve the required cleanliness. Please advise: the anticipated number of cycles, the total cleaning duration per campaign, and the specific manual interventions required at each step (drain valve operation, sample point access, lockout re-set, etc.).",
                "200-L 反应器进行单次\"喷淋-排空\"循环时，排空阶段远长于喷淋阶段，需多次循环以达到清洁度要求。请说明：预期循环次数、每次清洗总时长，以及各步骤所需的具体人工干预（排放阀操作 / 取样口操作 / 联锁复位等）。",
            ),
            (
                "Q1.9",
                "Cleaning sign-off authority and record",
                "清洗放行权限与记录",
                "Please confirm the role authorised to approve cleaning completion prior to the next campaign, and the documentation and records required to support this sign-off (batch record, cleaning log, sample analysis certificate, etc.).",
                "请确认在下一批次前有权审批清洗完成的岗位，以及支持该放行签批的文档与记录要求（批记录 / 清洗日志 / 样品分析证书等）。",
            ),
        ],
    },
    {
        "num_en": "2. 3D Layout & Equipment Space",
        "num_cn": "3D 布置与设备空间",
        "questions": [
            (
                "Q2.1",
                "Operational access clearance",
                "操作通道净空",
                "The current 3D model shows approximately 300–350 mm clearance in certain operator access zones. Please confirm the target minimum clearance from a site operability and maintenance standpoint, and identify any specific zones where relaxation of the current clearance is acceptable, or where additional clearance is required.",
                "当前 3D 模型显示部分操作人员通道净空约为 300–350 mm。请从现场操作与维护角度，确认目标最小净空要求，并识别可以放宽的具体区域，以及需要额外净空的区域。",
            ),
            (
                "Q2.2",
                "Free-drain feasibility from each vessel",
                "各容器的自重排放可行性",
                "The cleaning strategy relies on free-drain of cleaning fluid from each vessel into downstream collection (BT04 or equivalent). Considering the skid height constraint and each vessel elevation, please confirm the free-drain path is mechanically feasible from each source vessel to its destination. Please identify any vessels where drainage cannot be achieved by gravity and would require pumping or an alternative arrangement.",
                "清洗策略要求清洗液自重从各容器排入下游收集罐（BT04 或等效容器）。考虑撬块高度限制与各容器标高，请确认每台源容器至其目标容器的自重排放路径机械可行。请识别无法通过重力实现排放、需要泵抽或采用替代方案的容器。",
            ),
            (
                "Q2.3",
                "Nozzle orientation on pressure vessels",
                "压力容器管嘴方位",
                "Please review the nozzle orientation on all pressure vessels in the current 3D model. Confirm that nozzle positions provide adequate access for external piping tie-in, instrumentation installation, and maintenance (removal / reinstallation of internals, gasket replacement, catalyst service, etc.).",
                "请审查当前 3D 模型中所有压力容器的管嘴方位。确认管嘴位置能够满足外部管路对接、仪表安装以及维护（内构件拆装 / 密封垫更换 / 催化剂检维修等）的空间要求。",
            ),
            (
                "Q2.4",
                "Catalyst charging & discharging arrangement",
                "催化剂加料与卸料布置",
                "Please confirm that the catalyst charging (loading) and discharging (unloading) arrangement is compatible with the current nozzle layout on the reactor vessels, and that adequate space is provided for charge / discharge equipment and operator access during catalyst service.",
                "请确认催化剂加料与卸料的操作方式与当前反应柱的管嘴布置兼容，且在催化剂检维修期间为加料/卸料设备与操作人员留有充足空间。",
            ),
            (
                "Q2.5",
                "Tie-in points between skid and site utilities",
                "撬块与现场公用工程接口",
                "Please confirm the finalised location and specification of all tie-in points between the skid and the Sandwich site services, including process feed, product outlet, vent header, HTF supply and return, N₂ blanketing, instrument air, electrical supply, cable penetrations, and drain. Please advise the interfaces where design lock-down is still pending.",
                "请确认撬块与 Sandwich 现场公用工程之间所有接口的最终位置与规格，包括工艺进料、产品出口、放空总管、导热油供回、氮气钝化、仪表空气、电气供应、电缆穿墙、排放。请说明尚未定稿的接口。",
            ),
            (
                "Q2.6",
                "Cable wall penetration capacity and cable compatibility",
                "穿墙电缆孔容量与电缆兼容性",
                "The current design indicates 14 available cable wall penetration slots. Please confirm this final count, the internal diameter of each slot, and the compatibility of the slot with the largest cable type intended to pass through (multi-core, armoured, or shielded cables of relevant cross-section).",
                "当前设计显示可用穿墙电缆孔 14 个。请确认该数量的最终定案、每个孔的内径规格，以及与拟穿过的最大电缆类型（相应截面的多芯 / 铠装 / 屏蔽电缆）的兼容性。",
            ),
        ],
    },
    {
        "num_en": "3. Control Philosophy",
        "num_cn": "控制策略",
        "questions": [
            (
                "Q3.1",
                "Full Control Philosophy document issue date",
                "完整《控制策略》文件发出日期",
                "Please advise the target issue date for the complete Control Philosophy document, superseding the interim draft version previously circulated.",
                "请说明完整版《控制策略》文件的目标发出日期（覆盖此前流转的过渡草稿版本）。",
            ),
            (
                "Q3.2",
                "Interlock schedule — hardwired vs. software",
                "联锁清单 —— 硬件联锁与软件联锁",
                "Please provide a complete interlock schedule listing all safety interlocks, with clear delineation between hardwired (independent of the PLC) and software (PLC-implemented) interlocks. For each interlock, please include: initiating cause, sensing element, trip setpoint, action taken, and protective function fulfilled.",
                "请提供完整的联锁清单，列出全部安全联锁，明确区分硬件联锁（独立于 PLC）与软件联锁（PLC 实现）。每一条联锁请包含：触发原因、检测元件、跳停设定值、联锁动作、承担的保护功能。",
            ),
            (
                "Q3.3",
                "Operational modes and mode-specific fail positions",
                "操作模式与各模式下的故障位",
                "Please define all operational modes of the skid (normal operation, start-up, shut-down, cleaning, catalyst change, emergency depressurisation, etc.), and for each mode specify: (a) the fail position of each control valve; (b) the state of each interlock (active / bypassed); and (c) the permissive conditions required to enter that mode.",
                "请定义撬块的全部操作模式（正常运行 / 开车 / 停车 / 清洗 / 催化剂更换 / 紧急泄压等），并针对每一模式说明：(a) 每个控制阀的故障位；(b) 每个联锁的状态（激活 / 已绕过）；(c) 进入该模式所需的许可条件。",
            ),
            (
                "Q3.4",
                "Mode-transition safeguards",
                "操作模式切换的保护措施",
                "Please advise the safeguards in place against inadvertent mode transitions, particularly from cleaning mode back to processing mode (e.g. an interlock bypass being left in place after cleaning is complete). Please describe both the technical safeguards (key switches, HMI confirmations) and the procedural safeguards (SOP requirements, dual sign-off).",
                "请说明防止意外切换操作模式（尤其是从清洗模式返回处理模式时联锁绕过未复位）的保护措施。请分别描述技术性保护措施（钥匙开关 / HMI 确认）与程序性保护措施（SOP 要求 / 双签发）。",
            ),
        ],
    },
    {
        "num_en": "4. Pressure Vessel Mechanical Design",
        "num_cn": "压力容器机械设计",
        "questions": [
            (
                "Q4.1",
                "Design temperature vs. relief-scenario temperature",
                "设计温度 vs. 泄放工况温度",
                "The current equipment datasheet lists a design temperature of approximately 150 °C, whereas the relief-scenario analysis indicates process temperatures in the 250–300 °C range under credible upset conditions (fire case, cooling failure with reaction runaway, or combined scenarios). Please clarify: (a) whether the applicable pressure vessel code permits short-duration excursion above the design temperature during a relief event, and if so under what conditions; (b) whether the design temperature will be increased to 300 °C in the next revision of the datasheet; (c) the resulting impact on wall thickness and material allowable stress; and (d) any consequential changes to material selection, PWHT requirement, or hydrotest pressure.",
                "当前设备数据单标注的设计温度约 150 °C，而泄放工况分析显示在可预见异常工况下（火灾工况 / 冷媒失效叠加反应失控 / 组合工况）工艺温度可能达到 250–300 °C。请澄清：(a) 适用的压力容器规范是否允许在泄放事件中出现短时高温偏离，若允许则允许条件为何；(b) 下一版数据单是否将设计温度提高至 300 °C；(c) 由此对壁厚与材料许用应力的影响；(d) 由此可能带来的材料选型、PWHT 要求、水压试验压力等后续变化。",
            ),
            (
                "Q4.2",
                "Post-Weld Heat Treatment (PWHT) determination",
                "焊后热处理（PWHT）判定",
                "Based on the finalised wall thickness, material specification, and process service (hydrogen exposure), please confirm whether PWHT is mandatory under the applicable code for each pressure vessel in the skid. If PWHT is required, please advise the impact on the manufacturing schedule.",
                "根据最终确认的壁厚、材料规格与工艺工况（含氢环境），请确认撬块内各压力容器按适用规范是否强制要求 PWHT。若要求 PWHT，请说明对加工进度的影响。",
            ),
            (
                "Q4.3",
                "Design pressure & hydrotest pressure alignment",
                "设计压力与水压试验压力的一致性",
                "Please confirm the alignment of the design pressure (FV to 9 MPaG) and hydrotest pressure (1.5 × design pressure = 13.5 MPaG) across all pressure vessels in the skid, and update the equipment datasheets accordingly where discrepancies exist.",
                "请确认撬块内所有压力容器的设计压力（FV 至 9 MPaG）与水压试验压力（1.5 × 设计压力 = 13.5 MPaG）的一致性，如有差异请相应更新设备数据单。",
            ),
            (
                "Q4.4",
                "PED fluid group classification",
                "PED 介质类别分类",
                "Please confirm the PED fluid group classification for each pressure vessel and each associated pipe segment, particularly for services involving hydrogen and flammable organic solvents (which would typically fall under Group 1 per PED 2014/68/EU Annex II).",
                "请确认每台压力容器及关联管段的 PED 介质类别，特别是涉及氢气及易燃有机溶剂的工况（按 PED 2014/68/EU 附录 II 通常归入 Group 1）。",
            ),
            (
                "Q4.5",
                "Design calculation package issue date",
                "设计计算书发出日期",
                "Please confirm the target issue date for the complete design calculation package (shell strength, nozzle reinforcement, flange calculations, and FEA where applicable), which is a required component of the Technical File for CE conformity assessment.",
                "请确认完整设计计算书（壳体强度、管嘴补强、法兰计算、如适用的有限元分析）的目标发出日期，作为 CE 符合性评估技术文件的必备内容。",
            ),
        ],
    },
    {
        "num_en": "5. Sampler System",
        "num_cn": "取样系统",
        "questions": [
            (
                "Q5.1",
                "Sample vaporisation during pressure reduction",
                "取样过程中的样品闪蒸",
                "The sampler transfers process fluid from a line operating at up to 60 barG to a sample container at approximately atmospheric pressure. Please advise the analysis of vaporisation during this pressure reduction (fraction vaporised, temperature drop) and its impact on the compositional representativeness of the collected sample.",
                "取样器将最高 60 barG 工艺管路的介质转移至常压取样瓶。请说明此压降过程中的闪蒸分析（气化比例、温度降低）及其对所采样品成分代表性的影响。",
            ),
            (
                "Q5.2",
                "Sample retention design",
                "样品截留设计",
                "Please advise the design of the sample bottle and sample-transfer system, including any vapour trap, septum, condenser, or other retention device intended to prevent loss of volatiles through the vent path and to ensure a representative liquid sample is captured.",
                "请说明取样瓶及取样传输系统的设计，包括为防止挥发组分从放空路径流失、并确保获得代表性液样所采用的蒸汽阱、隔膜、冷凝器或其它截留装置。",
            ),
            (
                "Q5.3",
                "Sampler behaviour under gas-rich flow regime",
                "气相主导流态下取样器行为",
                "With the intended process operating at a stoichiometric excess of hydrogen, the flow regime through the sampling point may be gas-rich with entrained liquid rather than liquid-continuous. Please advise how the sampler design accommodates this flow regime, and whether any sample point relocation, phase separator, or coalescer is being considered.",
                "工艺预期在化学计量氢气过量条件下运行，取样点处流态可能为\"气相夹带液体\"而非液相连续。请说明取样器设计如何适应此流态，以及是否考虑取样点位置调整、加装相分离器或聚结器。",
            ),
        ],
    },
    {
        "num_en": "6. Instrumentation",
        "num_cn": "仪表配置",
        "questions": [
            (
                "Q6.1",
                "Guided-wave level transmitter — datasheet & cleanability",
                "Guided wave 液位变送器 —— 数据单与可清洁性",
                "Please provide the guided-wave level transmitter datasheet and confirm: (a) the process compatibility of the probe assembly with the intended service (including any operating conditions in the supercritical or near-critical regime); and (b) the cleanability of the probe assembly under the adopted cleaning philosophy.",
                "请提供 guided wave 液位变送器数据单，并确认：(a) 探头组件在拟定工况（含超临界或近临界工况）下的适用性；(b) 探头组件在拟采用清洗策略下的可清洁性。",
            ),
            (
                "Q6.2",
                "Level switch position & interlock function",
                "液位开关位置与联锁功能",
                "Please confirm the finalised position (elevation, orientation, flange size and rating) and the interlock function of each level switch (LSL, LSLL, LSH, LSHH) on each vessel, together with the matching instrument selection.",
                "请确认每台容器上各液位开关（低 LSL、低低 LSLL、高 LSH、高高 LSHH）的最终位置（标高、方位、法兰尺寸与压力等级）与联锁功能，以及对应的仪表选型。",
            ),
            (
                "Q6.3",
                "Instrument datasheet next revision date",
                "仪表数据单下一版本发出日期",
                "Please advise the target issue date for the next revision of the Instrument Datasheet, reflecting the current SE01 design (post demister-configuration decision), the finalised interlock schedule, and the level switch position confirmation.",
                "请说明反映当前 SE01 设计（demister 配置定稿后）、最终联锁清单以及液位开关位置确认的仪表数据单下一修订版本的目标发出日期。",
            ),
        ],
    },
    {
        "num_en": "7. Documentation Status & Pre-HAZOP Review Plan",
        "num_cn": "文档状态与 HAZOP 前评审安排",
        "questions": [
            (
                "Q7.1",
                "Current revision and next revision date of key deliverables",
                "主要交付文件的当前版本与下一版本时间",
                "Please confirm the current issued revision and the target date for the next revision of the following documents:  •  PFD  •  P&ID  •  URS  •  Equipment Datasheets  •  Control Philosophy  •  Instrument Datasheet  •  Operating Instructions  •  Cleaning Procedure  •  Vessel Design Calculation Package.",
                "请确认以下文档的当前已发版本与下一修订版本的目标日期：  ·  PFD  ·  P&ID  ·  URS  ·  设备数据单  ·  控制策略  ·  仪表数据单  ·  操作说明  ·  清洗程序  ·  容器设计计算书。",
            ),
            (
                "Q7.2",
                "Pre-HAZOP document review period",
                "HAZOP 前文档评审期",
                "Please advise the schedule for the pre-HAZOP document review period, during which the HAZOP study team will formally review the full document set prior to the HAZOP session. A minimum two-week review period ahead of the planned mid-August HAZOP is requested.",
                "请说明 HAZOP 前文档评审期的时间安排。HAZOP 分析小组将在本期内对完整文件集进行正式评审。请在 8 月中旬拟定 HAZOP 前预留至少两周的评审时间。",
            ),
            (
                "Q7.3",
                "HAZOP session logistics",
                "HAZOP 会议后勤安排",
                "Please confirm the target HAZOP session dates (mid-August 2026), the appointed HAZOP chair, the study team composition (site operations representative, process engineer, control & safety engineer, mechanical engineer, and IEPE / CFCT design representatives), the location, and any pre-work required from the design team.",
                "请确认 HAZOP 会议目标日期（2026 年 8 月中旬）、拟任 HAZOP 主席、分析小组构成（现场操作代表、工艺工程师、控制与安全工程师、机械工程师、以及 IEPE / CFCT 设计代表）、会议地点，以及设计小组会前需完成的准备工作。",
            ),
        ],
    },
]

for section in sections:
    add_section_header(doc, section["num_en"], section["num_cn"])
    for q in section["questions"]:
        add_question(doc, *q)

# Footer note
doc.add_paragraph()
tbl_end = doc.add_table(rows=1, cols=1)
tbl_end.autofit = False
tbl_end.columns[0].width = Cm(17)
cell_end = tbl_end.cell(0, 0)
cell_end.width = Cm(17)
set_cell_bg(cell_end, "FEF7ED")
tcPr = cell_end._tc.get_or_add_tcPr()
tcMar = OxmlElement('w:tcMar')
for side, val in [('top', 150), ('bottom', 150), ('left', 200), ('right', 200)]:
    node = OxmlElement(f'w:{side}')
    node.set(qn('w:w'), str(val))
    node.set(qn('w:type'), 'dxa')
    tcMar.append(node)
tcPr.append(tcMar)

p_end = cell_end.paragraphs[0]
p_end.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p_end.paragraph_format.line_spacing = 1.4
r_note = p_end.add_run("Note ｜ 说明：")
set_font(r_note, size=10, bold=True, color=ORANGE)

r_note_body = p_end.add_run(
    "This list represents outstanding design questions as of 2026-07-02. "
    "As alignment progresses, resolved items will be removed from "
    "subsequent revisions and any new items surfaced during document "
    "review or internal pre-HAZOP will be added.\n"
)
set_font(r_note_body, size=10, color=RGBColor(0x92, 0x40, 0x0E))

r_note_body_cn = p_end.add_run(
    "本清单为截至 2026-07-02 的待澄清设计问题。随对齐工作推进，已澄清的项将从后续版本中移除；"
    "如在文档评审或内部预 HAZOP 中出现新的问题，将补入清单。"
)
set_font(r_note_body_cn, size=9.5, color=RGBColor(0x92, 0x40, 0x0E))

# Save
out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Outstanding_Questions_for_Keith.docx",
)
doc.save(out_path)
print(f"Saved: {out_path}")
